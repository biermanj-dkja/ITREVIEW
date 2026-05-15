"""
report_generator_dg.py  —  DOCX report generator for Module 2
Data Governance and Data Flow Audit
v0.5.4.1

Report sections:
  1. Cover page
  2. Executive Summary — overall grade, top priorities, data-at-risk callout
  3. Per-system report cards — grade, per-area score bars, findings, strengths
  4. School-Wide Governance findings
  5. Action Plan — with Owner, Timing (Immediate / Near-Term / Planned), Effort
  6. Getting Started — 15-minute monthly ritual (shown for grade C and below)
  7. Appendix — raw answer log per system (with question prompt text)
"""

import io
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ────────────────────────────────────────────────
class C:
    urgent  = RGBColor(0xC0, 0x39, 0x2B)
    concern = RGBColor(0xE6, 0x7E, 0x22)
    watch   = RGBColor(0xF3, 0x9C, 0x12)
    healthy = RGBColor(0x27, 0xAE, 0x60)
    accent  = RGBColor(0x1A, 0x52, 0x76)
    mid     = RGBColor(0x2E, 0x86, 0xC1)
    text    = RGBColor(0x2C, 0x3E, 0x50)
    faint   = RGBColor(0x7F, 0x8C, 0x8D)
    white   = RGBColor(0xFF, 0xFF, 0xFF)
    silver  = RGBColor(0xBD, 0xC3, 0xC7)
    grade_a = RGBColor(0x1E, 0x8B, 0x4C)
    grade_b = RGBColor(0x1A, 0x52, 0x76)
    grade_c = RGBColor(0xB7, 0x77, 0x0D)
    grade_d = RGBColor(0xC0, 0x39, 0x2B)
    grade_f = RGBColor(0x8E, 0x1A, 0x1A)
    gs_bg   = RGBColor(0xEA, 0xF4, 0xFB)   # getting-started box background
    ritual  = RGBColor(0xE8, 0xF8, 0xF5)   # monthly ritual box background


SEV_COLOR = {
    "urgent":  C.urgent,
    "concern": C.concern,
    "watch":   C.watch,
    "healthy": C.healthy,
}
SEV_LABEL = {
    "urgent":  "URGENT",
    "concern": "CONCERN",
    "watch":   "WATCH",
    "healthy": "HEALTHY",
}
GRADE_COLOR = {
    "A": C.grade_a, "B": C.grade_b, "C": C.grade_c,
    "D": C.grade_d, "F": C.grade_f,
}
EFFORT_LABEL = {
    "S":  "Quick (½ day)",
    "S+": "Short (1 day)",
    "M":  "Medium (3 days)",
    "M+": "Substantial (5 days)",
    "L":  "Large (10 days)",
}
TIMING_LABEL = {
    "immediate":  "Do within 30 days",
    "near_term":  "Do within 90 days",
    "planned":    "Schedule this year",
}
TIMING_COLOR = {
    "immediate":  RGBColor(0xC0, 0x39, 0x2B),   # red
    "near_term":  RGBColor(0xE6, 0x7E, 0x22),   # orange
    "planned":    RGBColor(0x27, 0xAE, 0x60),   # green
}
AREA_ORDER = [
    "Access Control",
    "Backup & Recovery",
    "Data Flows",
    "Vendor & Contract",
    "Retention & Disposal",
    "School-Wide Governance",
]
SEV_ORDER = {"urgent": 0, "concern": 1, "watch": 2, "healthy": 3}
FSEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3}

# Question prompt text for the appendix (maps template_qid → short prompt)
QUESTION_PROMPTS = {
    "SYS.ID.name":   "System name",
    "SYS.ID.dept":   "Primary department",
    "SYS.ID.status": "System status",
    "SYS.ID.vendor": "Vendor / provider",
    "SYS.1.1":  "Can you generate a current list of every login account?",
    "SYS.1.1a": "Do any active accounts belong to former staff?",
    "SYS.1.2":  "Role-based access or flat access?",
    "SYS.1.2a": "Any staff with more admin access than needed?",
    "SYS.1.3":  "Is MFA enabled and required?",
    "SYS.1.4":  "Is this system connected to SSO?",
    "SYS.1.4a": "Shared or generic logins in use? (outside SSO)",
    "SYS.1.5":  "Does the system keep an audit log?",
    "SYS.2.1":  "Is the data backed up?",
    "SYS.2.1a": "Does the contract specify backup terms?",
    "SYS.2.1b": "Does the school maintain an independent copy?",
    "SYS.2.2":  "How frequently is the system backed up?",
    "SYS.2.3":  "When was the last restore test performed?",
    "SYS.2.4":  "Are backups stored separately from the live system?",
    "SYS.2.5":  "How long to restore if the system went down today?",
    "SYS.2.5v": "Documented interim process for vendor-hosted system outage?",
    "SYS.3.1":  "Where does data come from (inbound sources)?",
    "SYS.3.2":  "Where does data go (outbound destinations)?",
    "SYS.3.2a": "Are all outbound transfers encrypted?",
    "SYS.3.3":  "Does this system automatically exchange data with others?",
    "SYS.3.3a": "Details of automated connections",
    "SYS.3.4":  "Does your dept manually export and share data externally?",
    "SYS.3.4a": "Details of manual sharing",
    "SYS.3.5":  "Does the vendor sub-process data with other companies?",
    "SYS.4.1":  "Is there a signed Data Processing Agreement (DPA)?",
    "SYS.4.2":  "Does the contract require breach notification, and when?",
    "SYS.4.3":  "Is the vendor required to delete data when contract ends?",
    "SYS.4.4":  "Has the school reviewed vendor security practices?",
    "SYS.5.1":  "What types of data does this system hold?",
    "SYS.5.2":  "Required retention period for each data category",
    "SYS.5.3":  "How is data deleted at end of retention period?",
    "SYS.5.4":  "Decommissioned: was data exported and deletion confirmed?",
}


# ── Low-level helpers ─────────────────────────────────────────────

def _hex(c):
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def _cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_border(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _bottom_border(para, hex_color, sz=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _left_border(para, hex_color, sz=20):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _run(para, text, bold=False, italic=False, size=11, color=None):
    r = para.add_run(str(text or ""))
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or C.text
    return r


def _para(doc, text="", bold=False, italic=False, size=11, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def _h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 7)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18 if level == 1 else 13 if level == 2 else 11)
    r.font.color.rgb = C.accent if level < 3 else C.text
    if level == 1:
        _bottom_border(p, _hex(C.mid), sz=8)
    return p


def _box(doc, fill_hex, builder_fn):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, fill_hex)
    _cell_border(cell)
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    builder_fn(cell)
    if not cell.paragraphs:
        cell.add_paragraph()
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _cp(cell, text="", bold=False, italic=False, size=10, color=None, sb=2, sa=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def _set_hf(doc, school, report_date):
    sec = doc.sections[0]
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(hp, f"{school}  ·  Data Governance Audit Report", size=8, color=C.faint)
    _bottom_border(hp, _hex(C.silver), sz=4)
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)
    fp = ftr.add_paragraph()
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after = Pt(0)
    _run(fp, f"{report_date}    Page ", size=8, color=C.faint)
    r = fp.add_run()
    for tag in ("begin", "separate", "end"):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), tag)
        if tag == "begin":
            it = OxmlElement("w:instrText")
            it.text = "PAGE"
            r._r.append(fc)
            r._r.append(it)
        else:
            r._r.append(fc)
    r.font.size = Pt(8)
    r.font.color.rgb = C.faint


# ── Area score bar helper ─────────────────────────────────────────

def _area_score_table(doc, area_scores):
    """
    Render a compact per-area score breakdown table.
    area_scores: dict {area_name: (earned, max)}
    """
    if not area_scores:
        return
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, txt in enumerate(["Area", "Score", "Bar"]):
        c = tbl.rows[0].cells[i]
        _cell_bg(c, "EBF5FB")
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = C.accent

    for area in AREA_ORDER:
        if area not in area_scores:
            continue
        earned, max_pts = area_scores[area]
        pct = round(earned / max_pts * 100) if max_pts > 0 else 0
        row = tbl.add_row().cells
        _cell_bg(row[0], "FFFFFF")
        _cell_bg(row[1], "FFFFFF")
        _cell_bg(row[2], "FFFFFF")

        row[0].paragraphs[0].clear()
        r0 = row[0].paragraphs[0].add_run(area)
        r0.font.size = Pt(8)
        r0.font.color.rgb = C.text

        row[1].paragraphs[0].clear()
        bar_color = (C.healthy if pct >= 80 else C.watch if pct >= 60
                     else C.concern if pct >= 40 else C.urgent)
        r1 = row[1].paragraphs[0].add_run(f"{pct}%")
        r1.font.size = Pt(8)
        r1.bold = True
        r1.font.color.rgb = bar_color

        # Simple text bar using filled blocks
        filled = round(pct / 10)
        bar_str = "█" * filled + "░" * (10 - filled)
        row[2].paragraphs[0].clear()
        r2 = row[2].paragraphs[0].add_run(bar_str)
        r2.font.size = Pt(8)
        r2.font.color.rgb = bar_color

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Report sections ───────────────────────────────────────────────

def _cover(doc, school_name, respondent_name, respondent_role, report_date,
           is_draft=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(school_name)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = C.accent

    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(4)
    div.paragraph_format.space_after = Pt(20)
    _bottom_border(div, _hex(C.mid), sz=12)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after = Pt(8)
    tr = t.add_run("Data Governance and Data Flow Audit")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = C.text

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(24)
    _run(sub, "Findings Report and Action Plan", size=13, color=C.faint)

    if respondent_name or respondent_role:
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.paragraph_format.space_before = Pt(4)
        pf.paragraph_format.space_after = Pt(2)
        _run(pf, "Completed by", size=10, color=C.faint)
        if respondent_name:
            pn = doc.add_paragraph()
            pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pn.paragraph_format.space_before = Pt(2)
            pn.paragraph_format.space_after = Pt(2)
            _run(pn, respondent_name, bold=True, size=12)
        if respondent_role:
            pr = doc.add_paragraph()
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr.paragraph_format.space_before = Pt(0)
            pr.paragraph_format.space_after = Pt(4)
            _run(pr, respondent_role, size=10, color=C.faint)

    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_before = Pt(16)
    _run(pd, report_date, size=9, color=C.faint)

    if is_draft:
        draft_p = doc.add_paragraph()
        draft_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        draft_p.paragraph_format.space_before = Pt(24)
        draft_p.paragraph_format.space_after = Pt(4)
        dr = draft_p.add_run("DRAFT — ASSESSMENT INCOMPLETE")
        dr.bold = True
        dr.font.size = Pt(14)
        dr.font.color.rgb = C.concern
        draft_sub = doc.add_paragraph()
        draft_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        draft_sub.paragraph_format.space_before = Pt(0)
        _run(draft_sub,
             "Not all system worksheets or policy sections have been completed. "
             "Findings reflect only the systems assessed so far.",
             size=9, italic=True, color=C.faint)


def _dg_scope_box(doc, system_names, per_system_results):
    """
    Render an 'Assessment scope' callout box in the Executive Summary.
    system_names        — full list of systems the user registered
    per_system_results  — SystemResult list; only systems that were scored

    For each audited system, shows the data categories it holds (from
    SYS.5.1 answers) so a reader unfamiliar with the assessment knows
    exactly what was examined and why.
    """
    # Build a lookup: system_name -> SystemResult
    result_by_name = {r.system_name: r for r in per_system_results}
    audited_names  = [r.system_name for r in per_system_results]
    total_registered = len(system_names)
    total_audited    = len(audited_names)
    audited_set      = set(audited_names)
    not_audited      = [n for n in system_names if n not in audited_set]

    # What the DG audit actually examines — shown once at the top
    AUDIT_AREAS = [
        "data access controls and user permissions",
        "data retention and deletion practices",
        "vendor data-sharing and contractual protections",
        "backup coverage and recovery readiness",
        "breach notification and incident response procedures",
    ]

    def builder(cell):
        # Header
        p0 = _cp(cell, sb=3, sa=2)
        _run(p0, "Assessment Scope", bold=True, size=10, color=C.accent)
        if total_audited == total_registered:
            _run(p0, f"  -  all {total_registered} registered "
                     f"system{'s' if total_registered != 1 else ''} were audited.", size=10)
        else:
            _run(p0,
                 f"  -  {total_audited} of {total_registered} registered "
                 f"system{'s' if total_registered != 1 else ''} were audited in this session.",
                 size=10)

        # What the audit covers (static areas)
        pa = _cp(cell, sb=3, sa=1)
        _run(pa, "For each system, this audit reviewed: ", bold=True, size=9, color=C.faint)
        _run(pa, "; ".join(AUDIT_AREAS) + ".", size=9, color=C.faint)

        # Per-system breakdown with data categories
        if audited_names:
            ph = _cp(cell, sb=4, sa=1)
            _run(ph, "Systems audited in this session:", bold=True, size=9, color=C.accent)
            for name in audited_names:
                result = result_by_name.get(name)
                data_held = result.data_held if (result and result.data_held) else []
                p = _cp(cell, sb=1, sa=1)
                _run(p, f"  + {name}", bold=True, size=9, color=C.accent)
                if data_held:
                    _run(p, f"  (data held: {', '.join(data_held)})", size=9)
                else:
                    _run(p, "  (data categories not specified)", italic=True, size=9, color=C.faint)

        # Not audited
        if not_audited:
            ps = _cp(cell, sb=4, sa=1)
            _run(ps, "Registered but not audited in this session:", bold=True, size=9, color=C.faint)
            for name in not_audited:
                pp = _cp(cell, sb=1, sa=1)
                _run(pp, f"  - {name}", italic=True, size=9, color=C.faint)
            pn = _cp(cell, sb=3, sa=3)
            _run(pn,
                 "Findings and grades for unaudited systems are not included in this report.",
                 italic=True, size=9, color=C.faint)
        else:
            pe = _cp(cell, sb=4, sa=3)
            _run(pe, "All registered systems were audited - no gaps in coverage.",
                 italic=True, size=9, color=C.faint)

    _box(doc, "D6EAF8", builder)   # light blue, matches Module 1 scope box


def _exec_summary(doc, dg_report, school_name, system_names=None, is_draft=False):
    _page_break(doc)
    _h(doc, "Executive Summary", 1)

    # DRAFT callout — shown when assessment is not yet complete
    if is_draft:
        def draft_builder(cell):
            p = _cp(cell, sb=3, sa=2)
            _run(p, "⚠  DRAFT — Assessment Incomplete  ", bold=True, size=10, color=C.concern)
            _run(p, "Not all system worksheets or governance sections have been completed. "
                    "Grades, findings, and the action plan reflect only the systems assessed so far. "
                    "Complete remaining worksheets and regenerate this report for a full picture.",
                 size=10)
        _box(doc, "FDEDEC", draft_builder)

    # ── Assessment scope callout ─────────────────────────────────
    if system_names is not None:
        _dg_scope_box(doc, system_names, dg_report.per_system_results)

    summary = dg_report.summary

    # Overall stats box
    def stats_builder(cell):
        p1 = _cp(cell, sb=4, sa=2)
        _run(p1, "Overall Grade: ", bold=True, size=12)
        _run(p1, summary.overall_grade, bold=True, size=16,
             color=GRADE_COLOR.get(summary.overall_grade, C.text))
        _run(p1, f"   ·   {summary.total_systems} system{'' if summary.total_systems == 1 else 's'} audited",
             size=11, color=C.text)

        p2 = _cp(cell, sb=4, sa=2)
        if summary.critical_finding_count > 0:
            _run(p2, f"{summary.critical_finding_count} critical  ", bold=True, size=11, color=C.urgent)
        if summary.high_finding_count > 0:
            _run(p2, f"{summary.high_finding_count} high  ", bold=True, size=11, color=C.concern)

        p3 = _cp(cell, sb=4, sa=4)
        parts = []
        if summary.systems_urgent  > 0: parts.append(f"{summary.systems_urgent} urgent")
        if summary.systems_concern > 0: parts.append(f"{summary.systems_concern} concern")
        if summary.systems_watch   > 0: parts.append(f"{summary.systems_watch} watch")
        if summary.systems_healthy > 0: parts.append(f"{summary.systems_healthy} healthy")
        _run(p3, "Systems by severity: " + "  ·  ".join(parts) if parts else "No systems scored yet.",
             size=10, color=C.faint)

        # Data at risk callout
        if summary.data_at_risk_summary:
            p4 = _cp(cell, sb=6, sa=4)
            _run(p4, "⚠  ", bold=True, size=10, color=C.urgent)
            _run(p4, summary.data_at_risk_summary, size=10, color=C.text)

    fill = ("FDEDEC" if summary.overall_grade in ("D", "F") else
            "FEF9E7" if summary.overall_grade == "C" else "EAF4FB")
    _box(doc, fill, stats_builder)

    # Per-system grade table
    if dg_report.per_system_results:
        _h(doc, "System Report Card Summary", 2)
        _para(doc,
              "Scores reflect controls in place across access, backup, data flow, "
              "vendor agreements, and data retention for each system.",
              size=10, color=C.faint, sb=4, sa=8)

        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        for i, txt in enumerate(["System", "Grade", "Score", "Severity", "Findings"]):
            c = tbl.rows[0].cells[i]
            _cell_bg(c, _hex(C.accent))
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = C.white

        for idx, result in enumerate(dg_report.per_system_results):
            row = tbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill)

            grade_col = GRADE_COLOR.get(result.grade_label, C.text)
            sev_col = SEV_COLOR.get(result.severity, C.text)
            finding_count = len(result.findings)
            critical_n = sum(1 for f in result.findings if f.severity == "critical")
            finding_txt = (f"{finding_count} ({critical_n} critical)"
                           if critical_n else str(finding_count))

            for c, txt, col, bold in [
                (row[0], result.system_name,   C.text,    False),
                (row[1], result.grade_label,   grade_col, True),
                (row[2], f"{result.score_pct}%", C.text,  False),
                (row[3], SEV_LABEL.get(result.severity, result.severity), sev_col, True),
                (row[4], finding_txt, C.faint if finding_count == 0 else C.text, False),
            ]:
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(txt)
                r.font.size = Pt(9)
                r.font.color.rgb = col
                r.bold = bold

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Top Priorities box
    if summary.top_priorities:
        _h(doc, "Top Priorities", 2)
        _para(doc,
              "The highest-severity actions across all systems and school-wide findings, "
              "ordered by urgency. Assign an owner and deadline for each before this "
              "report is filed.",
              size=10, color=C.faint, sb=4, sa=8)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, txt in enumerate(["Finding", "System / Scope", "Owner", "Timing"]):
            c = tbl.rows[0].cells[i]
            _cell_bg(c, _hex(C.urgent))
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = C.white

        for idx, f in enumerate(summary.top_priorities):
            row = tbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "FEF5F5"
            for c in row:
                _cell_bg(c, fill)
            scope = f.system_name if f.system_name else "School-Wide"
            timing_txt = TIMING_LABEL.get(f.timing, f.timing)
            timing_col = TIMING_COLOR.get(f.timing, C.text)
            sev_col = SEV_COLOR.get(f.severity, C.text) if f.severity in ("critical",) else C.concern
            for c, txt, col, bold in [
                (row[0], f.title,     C.text,    False),
                (row[1], scope,       C.faint,   False),
                (row[2], f.owner,     C.text,    False),
                (row[3], timing_txt,  timing_col, True),
            ]:
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(txt)
                r.font.size = Pt(8)
                r.font.color.rgb = col
                r.bold = bold

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # School-wide summary line
    if summary.school_wide_findings:
        _h(doc, "School-Wide Governance", 2)
        _para(doc,
              f"{len(summary.school_wide_findings)} school-wide governance finding"
              f"{'s' if len(summary.school_wide_findings) != 1 else ''} identified — "
              "see the School-Wide Findings section for detail and recommended actions.",
              size=10, color=C.faint, sb=4, sa=8)


def _per_system_findings(doc, dg_report, finding_contexts=None):
    finding_contexts = finding_contexts or {}
    _page_break(doc)
    _h(doc, "Per-System Findings", 1)
    _para(doc,
          "One section per system. Findings are ordered by severity. "
          "Effort ratings: S=½ day · S+=1 day · M=3 days · M+=5 days · L=10 days.",
          size=10, color=C.faint, sb=4, sa=10)

    for result in dg_report.per_system_results:
        _h(doc, result.system_name, 2)

        # Score summary line
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(4)
        sp.paragraph_format.left_indent = Pt(6)
        grade_col = GRADE_COLOR.get(result.grade_label, C.text)
        sev_col = SEV_COLOR.get(result.severity, C.text)
        _run(sp, "Grade: ", size=10, color=C.faint)
        _run(sp, result.grade_label, bold=True, size=12, color=grade_col)
        _run(sp, f"   Score: {result.score_pct}% ({result.earned:.0f}/{result.max_pts} pts)",
             size=10, color=C.faint)
        _run(sp, "   Status: ", size=10, color=C.faint)
        _run(sp, SEV_LABEL.get(result.severity, result.severity),
             bold=True, size=10, color=sev_col)

        # Data categories held
        if result.data_held:
            from rules_engine_dg import _data_held_summary
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(2)
            dp.paragraph_format.space_after = Pt(6)
            dp.paragraph_format.left_indent = Pt(6)
            _run(dp, "Data held: ", size=9, color=C.faint)
            _run(dp, _data_held_summary(result.data_held), size=9, italic=True, color=C.text)

        # Per-area score breakdown
        if result.area_scores:
            _area_score_table(doc, result.area_scores)

        # No findings — show strengths
        if not result.findings:
            def healthy_builder(cell):
                p = _cp(cell, sb=4, sa=2)
                _run(p, "✓  All assessed controls appear to be in place.",
                     bold=True, size=10, color=C.healthy)
                if result.strengths:
                    for s in result.strengths:
                        ps = _cp(cell, sb=1, sa=1)
                        _run(ps, f"  ·  {s}", size=9, color=C.text)
            _box(doc, "EAFAF1", healthy_builder)
            continue

        # Findings sorted by severity then area
        sorted_findings = sorted(
            result.findings,
            key=lambda f: (FSEV_ORDER.get(f.severity, 9),
                           AREA_ORDER.index(f.area) if f.area in AREA_ORDER else 99)
        )

        for f in sorted_findings:
            sev_hex = _hex(SEV_COLOR.get(f.severity, C.text))
            # Build a stable finding_id key for context lookup
            fid_key = f"{result.section_id}:{f.area[:3].upper()}"

            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(8)
            tp.paragraph_format.space_after = Pt(3)
            tp.paragraph_format.left_indent = Pt(12)
            _left_border(tp, sev_hex, sz=16)
            _run(tp, f"[{f.severity.upper()}]  ", bold=True, size=9,
                 color=SEV_COLOR.get(f.severity, C.text))
            _run(tp, f.title, bold=True, size=10)
            _run(tp, f"  ·  {f.area}", size=8, color=C.faint)

            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(2)
            dp.paragraph_format.space_after = Pt(4)
            dp.paragraph_format.left_indent = Pt(12)
            _run(dp, f.detail, size=10)

            def action_builder(cell, f=f):
                p = _cp(cell, sb=3, sa=2)
                _run(p, "Action: ", bold=True, size=10, color=C.mid)
                _run(p, f.action, size=10)
                p2 = _cp(cell, sb=1, sa=3)
                timing_col = TIMING_COLOR.get(f.timing, C.faint)
                _run(p2, f"Owner: {f.owner}  ", italic=True, size=9, color=C.faint)
                _run(p2, f"  ·  {TIMING_LABEL.get(f.timing, f.timing)}",
                     italic=True, size=9, color=timing_col)
                if f.effort:
                    _run(p2, f"  ·  {EFFORT_LABEL.get(f.effort, f.effort)}",
                         italic=True, size=9, color=C.faint)

            _box(doc, "EBF5FB", action_builder)

            # Context note — shown when reviewer has annotated this finding
            ctx = finding_contexts.get(fid_key)
            if ctx:
                def ctx_builder(cell, ctx=ctx):
                    ph = _cp(cell, sb=2, sa=1)
                    _run(ph, "📋  Context note  ", bold=True, size=9, color=C.healthy)
                    _run(ph, f"(added {ctx['added_at'][:10]})", size=8, color=C.faint)
                    pb = _cp(cell, sb=1, sa=2)
                    _run(pb, ctx["note"], size=9, italic=True, color=C.text)
                _box(doc, "EAFAF1", ctx_builder)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _school_wide_findings(doc, dg_report):
    if not dg_report.school_wide_results:
        return
    _page_break(doc)
    _h(doc, "School-Wide Governance Findings", 1)
    _para(doc,
          "These findings apply across all systems rather than to any single platform. "
          "They represent policy, process, and governance gaps that affect the school's "
          "overall data protection posture.",
          size=10, color=C.faint, sb=4, sa=10)

    sorted_findings = sorted(
        dg_report.school_wide_results,
        key=lambda f: FSEV_ORDER.get(f.severity, 9)
    )

    for f in sorted_findings:
        sev_hex = _hex(SEV_COLOR.get(f.severity, C.text))

        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(10)
        tp.paragraph_format.space_after = Pt(3)
        tp.paragraph_format.left_indent = Pt(12)
        _left_border(tp, sev_hex, sz=16)
        _run(tp, f"[{f.severity.upper()}]  ", bold=True, size=9,
             color=SEV_COLOR.get(f.severity, C.text))
        _run(tp, f.title, bold=True, size=11)

        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(2)
        dp.paragraph_format.space_after = Pt(4)
        dp.paragraph_format.left_indent = Pt(12)
        _run(dp, f.detail, size=10)

        def action_builder(cell, f=f):
            p = _cp(cell, sb=3, sa=2)
            _run(p, "Action: ", bold=True, size=10, color=C.mid)
            _run(p, f.action, size=10)
            p2 = _cp(cell, sb=1, sa=3)
            timing_col = TIMING_COLOR.get(f.timing, C.faint)
            _run(p2, f"Owner: {f.owner}  ", italic=True, size=9, color=C.faint)
            _run(p2, f"  ·  {TIMING_LABEL.get(f.timing, f.timing)}",
                 italic=True, size=9, color=timing_col)
            if f.effort:
                _run(p2, f"  ·  {EFFORT_LABEL.get(f.effort, f.effort)}",
                     italic=True, size=9, color=C.faint)

        _box(doc, "EBF5FB", action_builder)


def _action_plan(doc, dg_report):
    _page_break(doc)
    _h(doc, "Action Plan", 1)
    _para(doc,
          "All recommended actions grouped by timing bucket, then by area. "
          "Each action lists the suggested owner role and estimated effort. "
          "Assign a named individual and a target date before this report is filed.",
          size=10, color=C.faint, sb=4, sa=10)

    # Collect all findings
    all_findings = []
    for result in dg_report.per_system_results:
        for f in result.findings:
            all_findings.append((f, result.system_name))
    for f in dg_report.school_wide_results:
        all_findings.append((f, "School-Wide"))

    # Split into timing buckets
    buckets = {
        "immediate":  [],
        "near_term":  [],
        "planned":    [],
    }
    for f, system_name in all_findings:
        bucket = getattr(f, "timing", "near_term") or "near_term"
        buckets[bucket].append((f, system_name))

    bucket_config = [
        ("immediate", "Immediate — Do Within 30 Days",
         "Critical and high-severity findings. Assign an owner and start date today.",
         "FDEDEC"),
        ("near_term", "Near-Term — Do Within 90 Days",
         "Medium-severity findings. Schedule these before the end of next quarter.",
         "FEF9E7"),
        ("planned",   "Planned — Schedule This Year",
         "Lower-severity findings and improvements. Add to the annual IT calendar.",
         "EAF4FB"),
    ]

    effort_rank = {"L": 0, "M+": 1, "M": 2, "S+": 3, "S": 4}
    fsev_rank = {"critical": 0, "high": 1, "medium": 2, "low": 3}

    for bucket_key, bucket_label, bucket_desc, bucket_fill in bucket_config:
        items = buckets[bucket_key]
        if not items:
            continue

        # Bucket header
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(12)
        ph.paragraph_format.space_after = Pt(2)
        rh = ph.add_run(bucket_label)
        rh.bold = True
        rh.font.size = Pt(12)
        timing_col = TIMING_COLOR.get(bucket_key, C.text)
        rh.font.color.rgb = timing_col
        _bottom_border(ph, _hex(timing_col), sz=4)

        _para(doc, bucket_desc, size=9, italic=True, color=C.faint, sb=2, sa=6)

        # Sort within bucket: severity first, then effort desc
        items_sorted = sorted(
            items,
            key=lambda x: (fsev_rank.get(x[0].severity, 9),
                           effort_rank.get(x[0].effort, 5))
        )

        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        for i, txt in enumerate(["Action", "System", "Owner", "Effort", "Severity"]):
            c = tbl.rows[0].cells[i]
            _cell_bg(c, bucket_fill)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = C.accent

        for idx, (f, system_name) in enumerate(items_sorted):
            row = tbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill)

            sev_col = SEV_COLOR.get(f.severity, C.text)
            effort_txt = EFFORT_LABEL.get(f.effort, f.effort or "—")

            for c, txt, col, bold in [
                (row[0], f.action,       C.text,  False),
                (row[1], system_name,    C.faint, False),
                (row[2], f.owner,        C.text,  False),
                (row[3], effort_txt,     C.text,  False),
                (row[4], f.severity.upper(), sev_col, True),
            ]:
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(str(txt))
                r.font.size = Pt(8)
                r.font.color.rgb = col
                r.bold = bold

        doc.add_paragraph().paragraph_format.space_after = Pt(10)


def _getting_started(doc, gs):
    """
    Render the Getting Started / Monthly Ritual section.
    Only shown for schools with grade C or below and significant gaps.
    Based on the Magic EdTech K-12 governance framework.
    """
    if not gs or not gs.show:
        return
    _page_break(doc)
    _h(doc, "Getting Started — Building a Governance Rhythm", 1)
    _para(doc,
          "The findings in this report can feel overwhelming when taken all at once. "
          "Research on K-12 data governance shows that the most effective approach is "
          "to start small: assign ownership, build one simple routine, and let it grow. "
          "The checklist and monthly ritual below are designed for schools at the "
          "beginning of that journey.",
          size=10, color=C.faint, sb=4, sa=10)

    # Five-step checklist
    _h(doc, "Your First Five Steps", 2)

    def checklist_builder(cell):
        for i, item in enumerate(gs.checklist, 1):
            p = _cp(cell, sb=3, sa=3)
            _run(p, f"{i}.  ", bold=True, size=10, color=C.accent)
            _run(p, item, size=10)

    _box(doc, _hex(C.gs_bg), checklist_builder)

    # Monthly ritual
    _h(doc, "The 15-Minute Monthly Governance Check-In", 2)
    _para(doc,
          "Schedule a recurring 15-minute meeting with IT and at least one member of "
          "leadership. The agenda is the same every month — keep it short, keep a "
          "written record, and rotate the host so it survives staff turnover.",
          size=10, color=C.faint, sb=4, sa=8)

    def ritual_builder(cell):
        ph = _cp(cell, sb=4, sa=2)
        _run(ph, "Monthly agenda:", bold=True, size=10, color=C.healthy)
        for item in gs.monthly_ritual_items:
            p = _cp(cell, sb=2, sa=2)
            _run(p, "  ✓  ", bold=True, size=9, color=C.healthy)
            _run(p, item, size=9)

    _box(doc, _hex(C.ritual), ritual_builder)


def _appendix(doc, dg_report, answers, system_names, generated_section_ids,
              amendment_log=None):
    _page_break(doc)
    _h(doc, "Appendix — Raw Response Log", 1)
    _para(doc,
          "Complete record of all answers submitted for each system worksheet, "
          "with question prompts for context. Questions with no answer are omitted.",
          size=10, color=C.faint, sb=4, sa=10)

    for i, (name, sid) in enumerate(zip(system_names, generated_section_ids), 1):
        _h(doc, f"{i}. {name}", 2)

        section_answers = {
            qid: data for qid, data in answers.items()
            if qid.startswith(f"{sid}_") and data.get("answer_status") not in ("unanswered", None)
        }

        if not section_answers:
            _para(doc, "No answers recorded for this system.", size=9, italic=True,
                  color=C.faint, sb=2, sa=6)
            continue

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i_h, txt in enumerate(["Question ID", "Question", "Status", "Answer"]):
            c = tbl.rows[0].cells[i_h]
            _cell_bg(c, "EBF5FB")
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = C.accent

        for idx, (qid, data) in enumerate(sorted(section_answers.items())):
            row = tbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill)

            raw = data.get("raw_answer")
            if isinstance(raw, list):
                answer_txt = ", ".join(str(x) for x in raw)
            else:
                answer_txt = str(raw) if raw else f"({data.get('answer_status', '—')})"
            answer_txt = answer_txt[:200]

            display_qid = qid.replace(f"{sid}_", "")
            prompt_txt = QUESTION_PROMPTS.get(display_qid, "")
            status = data.get("answer_status", "—")
            status_col = (C.concern if status == "unknown" else
                          C.faint if status == "skipped" else C.text)

            for c, txt, col in [
                (row[0], display_qid, C.faint),
                (row[1], prompt_txt,  C.faint),
                (row[2], status,      status_col),
                (row[3], answer_txt,  C.text),
            ]:
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(str(txt))
                r.font.size = Pt(8)
                r.font.color.rgb = col

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Answer Amendment Log — only if any revisions exist
    if amendment_log:
        _h(doc, "Answer Amendment Log", 2)
        _para(doc,
              "The following answers were changed after their section was initially marked complete.",
              size=10, color=C.faint, sb=4, sa=8)
        atbl = doc.add_table(rows=1, cols=4)
        atbl.style = "Table Grid"
        for i_h, txt in enumerate(["Q ID", "Changed At", "Previous Answer", "Revised Answer"]):
            c = atbl.rows[0].cells[i_h]
            _cell_bg(c, "EBF5FB")
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = C.accent

        for idx, row_data in enumerate(amendment_log):
            row = atbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill)
            old_val = str(row_data.get("old_raw_answer") or f"({row_data.get('old_answer_status', '—')})")[:120]
            new_val = str(row_data.get("new_raw_answer") or f"({row_data.get('new_answer_status', '—')})")[:120]
            changed = row_data.get("changed_at", "")[:16].replace("T", " ")
            for c, txt, col in [
                (row[0], row_data["question_id"], C.faint),
                (row[1], changed,                 C.faint),
                (row[2], old_val,                 C.concern),
                (row[3], new_val,                 C.text),
            ]:
                c.paragraphs[0].clear()
                run = c.paragraphs[0].add_run(str(txt))
                run.font.size = Pt(8)
                run.font.color.rgb = col


# ── Timeline helpers ───────────────────────────────────────────────

def _dg_sev_to_timeline(sev):
    """Map DG finding severity (critical/high/medium/low) → Module 1 scale (urgent/concern/watch)."""
    return {"critical": "urgent", "high": "concern", "medium": "watch", "low": "watch"}.get(sev, "watch")


def _dg_timing_to_horizon(timing):
    """Map DG timing (immediate/near_term/planned) → timeline horizon strings."""
    return {"immediate": "immediate", "near_term": "next_30_days", "planned": "next_90_days"}.get(timing, "next_30_days")


def _dg_timeline_section(doc, timeline):
    """Phased timeline section for the DG report — mirrors Module 1 _timeline_section."""
    _page_break(doc)
    _h(doc, "Phased Remediation Timeline", 1)
    _para(doc,
          "All recommended actions grouped by phase, ordered by severity then effort. "
          "Effort ratings: S=½ day · S+=1 day · M=3 days · M+=5 days · L=10 days.",
          size=10, sb=4, sa=10, color=C.faint)

    EFFORT_LABEL_TL = {
        "S":  "Quick (½ day)",
        "S+": "Short (1 day)",
        "M":  "Medium (3 days)",
        "M+": "Substantial (5 days)",
        "L":  "Large (10 days)",
    }
    SEV_COLOR_MAP = {"urgent": C.urgent, "concern": C.concern, "watch": C.watch}
    phase_fills   = {"1": "1A5276", "2": "1F618D", "3": "2471A3", "4": "2E86C1"}

    for phase in timeline["phases"]:
        acts = phase["actions"]
        if not acts and phase["phase"] in (2, 3, 4):
            continue

        # Phase header row
        fill_hex = phase_fills.get(str(phase["phase"]), "2C3E50")
        tbl_h = doc.add_table(rows=1, cols=1)
        tbl_h.style = "Table Grid"
        hcell = tbl_h.rows[0].cells[0]
        _cell_bg(hcell, fill_hex)
        _cell_border(hcell)
        for p in list(hcell.paragraphs):
            p._element.getparent().remove(p._element)
        hp = hcell.add_paragraph()
        hp.paragraph_format.space_before = Pt(4)
        hp.paragraph_format.space_after  = Pt(4)
        _run(hp, phase["label"], bold=True, size=12, color=C.white)
        if phase["duration_days"] > 0:
            date_str = f"  ·  {phase['start'].strftime('%d %b %Y')} → {phase['end'].strftime('%d %b %Y')}"
            _run(hp, date_str, size=10, color=C.white)
        else:
            _run(hp, f"  ·  Starting {phase['start'].strftime('%d %b %Y')}", size=10, color=C.white)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        if not acts:
            _para(doc, "No actions in this phase.", size=10, italic=True, color=C.faint, sb=2, sa=8)
        else:
            at = doc.add_table(rows=1, cols=4)
            at.style = "Table Grid"
            for i, txt in enumerate(["Action", "System / Area", "Effort", "Severity"]):
                c = at.rows[0].cells[i]
                _cell_bg(c, "EBF5FB")
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(txt)
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = C.accent

            effort_rank = {"L": 0, "M+": 1, "M": 2, "S+": 3, "S": 4}
            sev_rank    = {"urgent": 0, "concern": 1, "watch": 2}
            sorted_acts = sorted(acts, key=lambda a: (
                sev_rank.get(a.get("severity", "watch"), 9),
                effort_rank.get(a.get("effort", "M"), 5),
            ))

            for idx, act in enumerate(sorted_acts):
                row   = at.add_row().cells
                fill  = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
                for c in row:
                    _cell_bg(c, fill)
                sev_col    = SEV_COLOR_MAP.get(act.get("severity", "watch"), C.text)
                effort_txt = EFFORT_LABEL_TL.get(act.get("effort", ""), act.get("effort", "—"))
                # Use finding_title (system name) as the area column
                area_label = act.get("finding_title", act.get("section_id", "—"))

                for c, txt, col, bold_flag in [
                    (row[0], act["description"],                     C.text,  False),
                    (row[1], area_label,                             C.faint, False),
                    (row[2], effort_txt,                             C.text,  False),
                    (row[3], act.get("severity", "—").upper(),       sev_col, True),
                ]:
                    c.paragraphs[0].clear()
                    r = c.paragraphs[0].add_run(str(txt))
                    r.font.size = Pt(8)
                    r.font.color.rgb = col
                    r.bold = bold_flag

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        if phase.get("rerun_note"):
            def rnb(cell, note=phase["rerun_note"]):
                p = _cp(cell, sb=4, sa=4)
                _run(p, "⚡  ", bold=True, size=10, color=C.mid)
                _run(p, note, size=10, color=C.mid)
            _box(doc, "D6EAF8", rnb)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    strat = timeline.get("strategic_actions", [])
    if strat:
        _h(doc, "Strategic & Future Actions", 2)
        _para(doc,
              f"{len(strat)} action(s) are flagged as longer-horizon initiatives. "
              "Not included in the phased timeline — revisit annually.",
              size=10, sb=4, sa=6, color=C.faint)
        for act in strat:
            p = _para(doc, sb=2, sa=2)
            p.paragraph_format.left_indent = Pt(18)
            _run(p, act["description"], size=10)


# ── Main entry point ──────────────────────────────────────────────

def generate_dg_report(dg_report_obj, answers, profile, system_names,
                       generated_section_ids, start_date=None,
                       is_complete=True, finding_contexts=None, amendment_log=None):
    """
    Parameters
    ----------
    dg_report_obj        : DGReport from rules_engine_dg.evaluate_dg()
    answers              : dict from database.get_answers()
    profile              : dict from database.get_school_profile()
    system_names         : list of str
    generated_section_ids: list of str
    start_date           : ISO date string or None — if provided, a phased
                           remediation timeline section is added to the report.
    is_complete          : bool — False adds a DRAFT watermark to cover and exec summary
    finding_contexts     : dict of {finding_id: {note, added_at}} — reviewer annotations
    amendment_log        : list of amendment history dicts — shown in appendix if present

    Returns
    -------
    bytes — the DOCX file content
    """
    school_name = (profile or {}).get("school_name", "School")
    report_date = date.today().isoformat()
    is_draft = not is_complete
    finding_contexts = finding_contexts or {}

    def _get(qid):
        d = answers.get(qid)
        if not d:
            return ""
        r = d.get("raw_answer")
        return ", ".join(str(x) for x in r) if isinstance(r, list) else (str(r) if r else "")

    respondent_name = _get("DG1.1")
    respondent_role = _get("DG1.1b") if "DG1.1b" in answers else ""

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

    _set_hf(doc, school_name, report_date)
    _cover(doc, school_name, respondent_name, respondent_role, report_date,
           is_draft=is_draft)
    _exec_summary(doc, dg_report_obj, school_name, system_names=system_names,
                  is_draft=is_draft)
    _per_system_findings(doc, dg_report_obj, finding_contexts=finding_contexts)
    _school_wide_findings(doc, dg_report_obj)
    _action_plan(doc, dg_report_obj)

    # Phased remediation timeline (only when start_date supplied)
    if start_date:
        from timeline import build_timeline
        from datetime import date as _date
        sd = _date.fromisoformat(start_date) if isinstance(start_date, str) else start_date
        # Flatten DG findings into the same action-dict format build_timeline expects
        flat_findings = []
        for result in dg_report_obj.per_system_results:
            for f in result.findings:
                flat_findings.append({
                    "finding_id":    f"{result.section_id}:{f.area[:3].upper()}",
                    "title":         f.title,
                    "severity":      _dg_sev_to_timeline(f.severity),
                    "section_id":    result.section_id,
                    "actions": [{
                        "description":   f.action,
                        "effort":        f.effort or "M",
                        "time_horizon":  _dg_timing_to_horizon(getattr(f, "timing", "near_term")),
                        "constraint_flag": False,
                    }],
                })
        for f in dg_report_obj.school_wide_results:
            flat_findings.append({
                "finding_id":    f"DG2:{f.area[:3].upper()}",
                "title":         f.title,
                "severity":      _dg_sev_to_timeline(f.severity),
                "section_id":    "DG2",
                "actions": [{
                    "description":   f.action,
                    "effort":        f.effort or "M",
                    "time_horizon":  _dg_timing_to_horizon(getattr(f, "timing", "near_term")),
                    "constraint_flag": False,
                }],
            })
        if flat_findings:
            timeline = build_timeline(flat_findings, sd)
            _dg_timeline_section(doc, timeline)

    _getting_started(doc, dg_report_obj.getting_started)
    _appendix(doc, dg_report_obj, answers, system_names, generated_section_ids,
              amendment_log=amendment_log)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
