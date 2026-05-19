"""
Report Generator for Module 1 — School IT State of the System Report
v0.5.4.1

Uses python-docx to produce the DOCX entirely in Python.

Report structure:
  1. Cover page
  2. Table of Contents (native Word TOC field — updates on open)
  3. Executive Summary
       - Overall health verdict (one sentence)
       - Section Scores table (with What's Working column)
       - Data confidence callout (in body, not just footer)
       - Priority Findings boxes
  4. Key Risks  (primary finding labelled "Start here")
  5. Section-by-Section Findings
       (healthy sections get a ✓ marker; removed from appendix-only)
  6. Phased Remediation Timeline  ← ACTION PLAN (bullets removed; table only)
  7. Appendix
       A. Composite Finding Traceability
       B. Unknown Answer Log  (prominently framed, moved before C)
       C. Full Response Log  (with question text column)
"""

import io
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────
class C:
    urgent  = RGBColor(0xC0, 0x39, 0x2B)
    concern = RGBColor(0xE6, 0x7E, 0x22)
    watch   = RGBColor(0xF3, 0x9C, 0x12)
    healthy = RGBColor(0x27, 0xAE, 0x60)
    accent  = RGBColor(0x1A, 0x52, 0x76)
    mid     = RGBColor(0x2E, 0x86, 0xC1)
    text    = RGBColor(0x2C, 0x3E, 0x50)
    faint   = RGBColor(0x7F, 0x8C, 0x8D)
    white   = RGBColor(0xFF, 0xFF, 0xFF)
    silver  = RGBColor(0xBD, 0xC3, 0xC7)

SEV_COLOR  = {"urgent": C.urgent, "concern": C.concern, "watch": C.watch, "healthy": C.healthy}
SEV_LABEL  = {"urgent": "URGENT", "concern": "CONCERN", "watch": "WATCH",
               "healthy": "HEALTHY", "context_only": "CONTEXT ONLY"}
HORIZON_LABEL = {
    "immediate":       "Immediate",
    "next_30_days":    "Within 30 days",
    "next_90_days":    "Within 90 days",
    "next_12_months":  "Within 12 months",
    "strategic_future": "Strategic / Future",
}
HORIZON_ORDER = ["immediate", "next_30_days", "next_90_days", "next_12_months", "strategic_future"]
SECTION_NAMES = {
    "2": "Governance, Budget, Staffing, and Ownership",
    "3": "Sites, Buildings, Network, and Internet",
    "4": "Identity, Accounts, and Access",
    "5": "Endpoints, Printing, and Classroom Technology",
    "6": "Core Systems, Servers, Vendors, and Contracts",
    "7": "Data Protection, Backup, and Recovery",
    "8": "Security Operations, Filtering, and Safeguards",
    "9": "Documentation and Operational Readiness",
}
# One-sentence plain-English description of what each section examines.
# Used in the dynamic scope statement so a reader who did not fill out
# the assessment understands exactly what was (and was not) reviewed.
SECTION_DESCRIPTIONS = {
    "2": "how IT is staffed, funded, planned, and managed — including budget, vendor contracts, and written policies",
    "3": "the physical network infrastructure: wiring, wireless access, firewalls, and internet service across all buildings and sites",
    "4": "user accounts, passwords, directory services, multi-factor authentication, and who has access to what",
    "5": "staff and student computers, laptops, tablets, printers, projectors, and classroom A/V equipment",
    "6": "core business systems (SIS, LMS, email, finance), on-premise servers, and key vendor relationships",
    "7": "how school data is backed up, how quickly it can be restored, and whether recovery has ever been tested",
    "8": "security monitoring, web filtering, anti-malware, phishing defenses, and student safeguarding controls",
    "9": "IT runbooks, asset inventories, onboarding/offboarding procedures, and operational continuity documentation",
}

SEV_ORDER  = {"urgent": 0, "concern": 1, "watch": 2, "healthy": 3, "context_only": 4}
FSEV_ORDER = {"urgent": 0, "concern": 1, "watch": 2}

# Question text for the full response log (section.question → short prompt)
QUESTION_PROMPTS = {
    "1.1":  "School name",
    "1.2":  "School address",
    "1.3":  "Phone number",
    "1.4":  "Domain / website",
    "1.5":  "School mission / tagline",
    "1.6":  "Accredited?",
    "1.7a": "Respondent name",
    "1.7b": "Respondent role",
    "1.8":  "Number of campuses / sites",
    "1.9":  "Number of buildings (total across all sites)",
    "1.10": "Buildings per site",
    "1.11": "Grades served",
    "1.12": "Approximate student enrollment",
    "1.13": "Approximate faculty and staff count",
    "1.14": "Approximate total managed devices",
    "1.15": "Device categories in scope",
    "1.16": "Upcoming calendar events affecting IT planning",
    "2.1":  "IT staffing model",
    "2.2":  "Dedicated IT budget?",
    "2.3":  "IT budget reviewed annually?",
    "2.4":  "IT spending tracked?",
    "2.5":  "Technology plan or roadmap?",
    "2.6":  "Vendor management process?",
    "2.7":  "Contracts and renewals tracked?",
    "2.8":  "IT policies documented?",
    "2.9":  "IT coverage if primary person unavailable?",
    "2.10": "Recurring IT tasks tracked?",
    "2.11": "Technology committee or advisory group?",
    "2.12": "Ed tech / academic technology role?",
    "3.1":  "Network diagram — current?",
    "3.2":  "Wiring diagrams — current?",
    "3.3":  "ISP circuits documented?",
    "3.4":  "Wireless AP locations documented?",
    "3.5":  "Wireless coverage survey done?",
    "3.6":  "Switch inventory current?",
    "3.7":  "VLANs / network segmentation?",
    "3.8":  "Known connectivity issues?",
    "3.9":  "Router inventory current?",
    "3.10": "Firewall inventory current?",
    "3.11": "Firewall under support contract?",
    "3.12": "Admin access to all network infrastructure?",
    "3.13": "Number of ISP circuits",
    "3.14": "ISP providers",
    "3.15": "ISP failover / redundancy?",
    "3.16": "Guest network separated?",
    "3.18": "Network device configs backed up?",
    "3.19": "Wireless coverage adequate?",
    "3.20": "Network documentation completeness?",
    "3.21": "IoT / classroom devices segmented?",
    "3.23": "Network monitoring in place?",
    "3.24": "Vulnerability scanning in place?",
    "3.26": "Known unresolved security concerns?",
    "4.1":  "Identity platform (Google, Microsoft, hybrid)?",
    "4.1b": "SSO in use?",
    "4.2":  "Directory service?",
    "4.3":  "Onboarding process documented?",
    "4.4":  "Offboarding process documented?",
    "4.5":  "Privilege review process?",
    "4.6":  "MFA on privileged accounts?",
    "4.6b": "MFA required for all staff?",
    "4.7":  "Access audit completed within 12 months?",
    "4.8":  "Shared accounts documented?",
    "4.9":  "Admin credential documentation?",
    "5.1":  "Device inventory current?",
    "5.3":  "BYOD policy?",
    "5.5":  "Device management (MDM) in use?",
    "5.6":  "Hardware standard documented?",
    "5.7":  "Software standard documented?",
    "5.8":  "Remote support capability?",
    "5.9":  "Device refresh cycle documented?",
    "5.10": "End-of-life devices in use?",
    "5.11": "Warranty tracking?",
    "5.12": "Loaner / spare device process?",
    "5.13": "Device check-in/check-out process?",
    "5.14": "Average device age (years)",
    "5.15": "Student devices managed centrally?",
    "5.17": "Printer inventory tracked?",
    "5.18": "Number of printers",
    "5.19": "Average printer age (years)",
    "6.1":  "Core systems list",
    "6.2":  "Systems list location / format",
    "6.3":  "Systems list current?",
    "6.5":  "Known system dependencies",
    "6.6":  "Contract renewal tracking?",
    "6.7":  "Privacy / DPA review for student-data systems?",
    "6.8":  "Vendor contact list current?",
    "6.9":  "Escalation paths documented for critical vendors?",
    "6.10": "Servers in use?",
    "6.11": "Server OS patching cycle defined?",
    "6.12": "Server details / notes",
    "6.13": "Number of on-premises servers",
    "7.1":  "Backup policy documented?",
    "7.2":  "Backup systems / tools in use",
    "7.3":  "Backup scope documented?",
    "7.4":  "Server backup status",
    "7.5":  "Staff device backups?",
    "7.6":  "Critical cloud data backed up?",
    "7.7":  "Backup health reviewed regularly?",
    "7.7b": "Backup storage location (onsite / offsite / both)?",
    "7.8":  "Restore test completed within 12 months?",
    "7.10": "Recovery priority documented?",
    "7.11": "RTO / RPO reference documented?",
    "7.12": "Emergency credential access?",
    "7.13": "Estimated full recovery time",
    "8.1":  "Endpoint protection deployed?",
    "8.3":  "Patch management schedule?",
    "8.4":  "Security policy reviewed in last 12 months?",
    "8.5":  "Content filtering in use?",
    "8.6":  "Email security / anti-phishing?",
    "8.7":  "Data loss prevention controls documented?",
    "8.8":  "Incident response process?",
    "8.9":  "Security logs reviewed regularly?",
    "8.10": "Staff security awareness training?",
    "8.11": "Known unresolved security concerns?",
    "9.1":  "Documentation system in use?",
    "9.2":  "Documentation currency?",
    "9.3":  "SOPs documented for recurring tasks?",
    "9.4":  "Change documentation process?",
    "9.5":  "Environment understandable from documentation alone?",
    "9.6":  "Knowledge concentration / single-person dependency?",
    "10.1": "Top IT priorities for next 12 months",
    "10.2": "Upcoming projects or initiatives",
    "10.3": "Known risks or concerns",
    "10.4": "Budget or planning notes",
    "10.5": "Planning to make any major changes?",
    "10.6": "Major change details",
    "10.7": "Confidence in assessment answers",
    "10.8": "Anything else to add",
}


# ── Low-level helpers ─────────────────────────────────────────────

def _hex(c):
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def _cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_border(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _bottom_border(para, hex_color, sz=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _left_border(para, hex_color, sz=20):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _run(para, text, bold=False, italic=False, size=11, color=None):
    r = para.add_run(str(text or ""))
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or C.text
    return r


def _para(doc, text="", bold=False, italic=False, size=11, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def _h(doc, text, level=1):
    """
    Heading using a named Word style (Heading 1/2/3) so the TOC field
    can pick them up. Falls back to plain paragraph if the style is absent.
    The visual appearance is set via the style definitions in _apply_styles().
    We add a bottom border only on H1 for visual emphasis.
    """
    style_name = f"Heading {level}"
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()

    p.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 7)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18 if level == 1 else 13 if level == 2 else 11)
    r.font.color.rgb = C.accent if level < 3 else C.text
    if level == 1:
        _bottom_border(p, _hex(C.mid), sz=8)
    return p


def _box(doc, fill_hex, builder_fn):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _cell_bg(cell, fill_hex)
    _cell_border(cell)
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    builder_fn(cell)
    if not cell.paragraphs:
        cell.add_paragraph()
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _cp(cell, text="", bold=False, italic=False, size=10, color=None, sb=2, sa=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def _get(answers, qid):
    d = answers.get(qid)
    if not d:
        return ""
    r = d.get("raw_answer")
    return ", ".join(str(x) for x in r) if isinstance(r, list) else (str(r) if r else "")


def _cost_tier(actions):
    """
    Derive a budget signal from a finding's action list.
    Returns (label, hex_color) based on the highest-effort action and
    whether any constraint_flag is set.
    Tiers:
      Staff time only  — all actions S or S+, no purchase keywords
      Moderate budget  — at least one M or M+ action, or purchase keyword but no L
      Significant investment — at least one L action, or constraint_flag on any action
    """
    if not actions:
        return None, None
    purchase_kw = ("purchase", "procure", "buy", "invest", "replac", "licens", "subscri", "vendor")
    has_L        = any(a.get("effort") == "L"  for a in actions)
    has_M        = any(a.get("effort") in ("M", "M+") for a in actions)
    has_purchase = any(any(kw in (a.get("description") or "").lower() for kw in purchase_kw)
                       for a in actions)
    has_constraint = any(a.get("constraint_flag") for a in actions)

    if has_L or has_constraint:
        return "Significant investment", "C0392B"   # red-ish
    elif has_M or has_purchase:
        return "Moderate budget",       "E67E22"   # orange
    else:
        return "Staff time only",       "1A7A4A"   # green


def _scope_box(doc, sections_covered, skipped_section_ids):
    """
    Render a light-blue 'Assessment scope' callout near the top of the
    Executive Summary.  Lists each covered section by name and a plain-English
    description of its subject matter so a reader who did not complete the
    questionnaire can quickly understand what was (and was not) reviewed.
    """
    total = len(SECTION_NAMES)
    covered_ids = sorted(set(SECTION_NAMES.keys()) - set(skipped_section_ids))
    skipped_ids_sorted = sorted(skipped_section_ids)

    def builder(cell):
        # Header
        p0 = _cp(cell, sb=3, sa=3)
        _run(p0, "Assessment Scope", bold=True, size=10, color=C.accent)
        if sections_covered == total:
            _run(p0, f"  -  all {total} IT domains were reviewed in this assessment.", size=10)
        else:
            _run(p0,
                 f"  -  {sections_covered} of {total} IT domains were reviewed "
                 f"in this assessment.",
                 size=10)

        # Covered sections
        for sid in covered_ids:
            name = SECTION_NAMES.get(sid, f"Section {sid}")
            desc = SECTION_DESCRIPTIONS.get(sid, "")
            p = _cp(cell, sb=1, sa=1)
            _run(p, f"  + {name}", bold=True, size=9, color=C.accent)
            if desc:
                _run(p, f":  {desc}.", size=9)

        # Skipped sections
        if skipped_ids_sorted:
            ps = _cp(cell, sb=4, sa=1)
            _run(ps, "Not reviewed in this session:", bold=True, size=9, color=C.faint)
            for sid in skipped_ids_sorted:
                name = SECTION_NAMES.get(sid, f"Section {sid}")
                desc = SECTION_DESCRIPTIONS.get(sid, "")
                pp = _cp(cell, sb=1, sa=1)
                _run(pp, f"  - {name}", bold=True, size=9, color=C.faint)
                if desc:
                    _run(pp, f":  {desc}.", italic=True, size=9, color=C.faint)
            pn = _cp(cell, sb=3, sa=3)
            _run(pn,
                 "Findings and scores for unreviewed domains are not included in this report.",
                 italic=True, size=9, color=C.faint)
        else:
            pe = _cp(cell, sb=4, sa=3)
            _run(pe, "All IT domains were assessed - no gaps in coverage.", italic=True, size=9, color=C.faint)

    _box(doc, "D6EAF8", builder)   # light blue


def _set_hf(doc, school, report_date, caveat, assessment_date=None):
    sec = doc.sections[0]
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(hp, f"{school}  ·  IT State of the System Report", size=8, color=C.faint)
    _bottom_border(hp, _hex(C.silver), sz=4)

    ftr = sec.footer
    ftr.is_linked_to_previous = False
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)
    if caveat:
        cp = ftr.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(2)
        _run(cp, f"⚠ Data confidence: {caveat}", italic=True, size=7, color=C.concern)
    fp = ftr.add_paragraph()
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after = Pt(0)
    date_label = f"Assessed: {assessment_date}  ·  Generated: {report_date}" \
                 if assessment_date and assessment_date != report_date else report_date
    _run(fp, f"{date_label}    Page ", size=8, color=C.faint)
    r = fp.add_run()
    for tag in ("begin", "separate", "end"):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), tag)
        if tag == "begin":
            it = OxmlElement("w:instrText")
            it.text = "PAGE"
            r._r.append(fc)
            r._r.append(it)
        else:
            r._r.append(fc)
    r.font.size = Pt(8)
    r.font.color.rgb = C.faint


def _apply_heading_styles(doc):
    """
    Ensure Heading 1/2/3 styles exist with outlineLevel set so the TOC
    field can collect them. python-docx creates these lazily; touching
    them here guarantees they're in styles.xml before content is added.
    """
    from docx.oxml.ns import nsmap
    styles = doc.styles
    level_map = {
        "Heading 1": (0, 18, True),
        "Heading 2": (1, 13, True),
        "Heading 3": (2, 11, True),
    }
    for style_name, (outline_level, pt_size, bold) in level_map.items():
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, 1)  # 1 = paragraph style
        # Set outlineLevel via XML so Word includes these in the TOC
        pPr = style.element.get_or_add_pPr()
        existing_ol = pPr.find(qn("w:outlineLvl"))
        if existing_ol is None:
            ol = OxmlElement("w:outlineLvl")
            ol.set(qn("w:val"), str(outline_level))
            pPr.append(ol)
        else:
            existing_ol.set(qn("w:val"), str(outline_level))


# ── Table of Contents ─────────────────────────────────────────────

def _toc(doc):
    """
    Insert a native Word TOC field. Word/LibreOffice will render it
    when the document is opened; right-click → Update Field to refresh.
    Uses Heading 1 and Heading 2 (TOC \o "1-2").
    """
    _page_break(doc)
    _h(doc, "Contents", 1)

    # TOC field paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "false")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "[Right-click here and choose 'Update Field' to generate the table of contents]"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    run.font.size = Pt(10)
    run.font.color.rgb = C.faint


# ── Overall health verdict ────────────────────────────────────────

def _health_verdict(summary, scores):
    """
    Return a one-sentence plain-language health statement.
    """
    urgent = summary["urgent_count"]
    concern = summary["concern_count"]
    # Count healthy / watch sections
    healthy_secs = [s for s in scores if s.get("severity") == "healthy"]
    concern_secs = [s for s in scores if s.get("severity") in ("concern", "urgent")]

    if urgent >= 3:
        return (
            f"The assessment found {urgent} urgent findings that require immediate action — "
            "the school's IT environment has significant gaps that create real risk today."
        )
    elif urgent >= 1:
        return (
            f"The assessment found {urgent} urgent and {concern} concern-level findings. "
            "Several areas need prompt attention, though a number of sections are performing well."
        )
    elif concern >= 4:
        return (
            f"No urgent findings, but {concern} concern-level findings indicate meaningful gaps "
            "across several areas. The environment is functional but needs systematic improvement."
        )
    elif concern >= 1:
        return (
            f"The assessment found {concern} concern-level findings in specific areas. "
            f"{len(healthy_secs)} of {len(scores)} scored sections are healthy — "
            "the fundamentals are largely in place."
        )
    else:
        return (
            "The assessment found no urgent or concern-level findings. "
            "The IT environment is well-managed — focus on watch-level items to maintain momentum."
        )


# ── Staffing stub ─────────────────────────────────────────────────

def _staffing_stub(doc, answers):
    """
    Render a short narrative paragraph about the IT staffing model drawn
    directly from Section 2 answers.  Used in the Executive Summary.
    Reads: 2.1 (model), 2.2 (named lead), 2.2b (lead name/title),
           2.3 (day-to-day assigned), 2.7 (known constraints),
           2.9 (continuity).
    """
    model       = _get(answers, "2.1")
    has_lead    = _get(answers, "2.2")
    lead_name   = _get(answers, "2.2b")
    day_to_day  = _get(answers, "2.3")
    constraints = _get(answers, "2.7")
    continuity  = _get(answers, "2.9")

    # Skip entirely if we have almost nothing to say
    if not any([model, has_lead, lead_name]):
        return

    # ── Build sentence fragments ──────────────────────────────────
    if model:
        model_clean = model.strip().rstrip(".")
        model_sent = f"IT support is currently structured as: {model_clean}."
    else:
        model_sent = "The IT staffing model was not specified."

    if has_lead == "yes" and lead_name:
        lead_sent = f"{lead_name.strip()} serves as the named IT lead."
    elif has_lead == "yes":
        lead_sent = "A named IT lead is in place."
    elif has_lead == "no":
        lead_sent = "No single named person is currently accountable for IT leadership."
    else:
        lead_sent = ""

    if day_to_day == "yes":
        dod_sent = "Day-to-day support responsibilities are clearly assigned."
    elif day_to_day == "no":
        dod_sent = "Day-to-day IT support responsibilities are not formally assigned."
    else:
        dod_sent = ""

    continuity_map = {
        "Yes — another person could cover fully":
            "If the primary lead were unavailable, another person could cover operations.",
        "Partially — some things would be managed, others would not":
            "Partial coverage exists if the primary lead were unavailable; some tasks would go unmanaged.",
        "No — operations would be significantly disrupted":
            "There is currently no continuity plan — operations would be significantly disrupted "
            "if the IT lead were unavailable.",
    }
    cont_sent = continuity_map.get(continuity, "")

    if constraints and constraints.lower() not in ("no", "n/a", "none", ""):
        constraint_sent = (
            f"Known constraint: {constraints.strip().rstrip('.')}."
            if constraints.lower() not in ("yes",)
            else "Known budget or staffing constraints are noted for this year."
        )
    else:
        constraint_sent = ""

    # ── Render ────────────────────────────────────────────────────
    _h(doc, "IT Staffing Overview", 2)

    def builder(cell):
        sentences = [s for s in [model_sent, lead_sent, dod_sent, cont_sent] if s]
        p = _cp(cell, sb=4, sa=4)
        _run(p, "  ".join(sentences), size=10)
        if constraint_sent:
            p2 = _cp(cell, sb=2, sa=4)
            _run(p2, "\u26a0  ", bold=True, size=9, color=C.concern)
            _run(p2, constraint_sent, italic=True, size=9, color=C.concern)

    _box(doc, "EBF5FB", builder)


# ── Cover ─────────────────────────────────────────────────────────

def _cover(doc, meta):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(meta.get("school_name") or "School")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = C.accent

    if meta.get("school_mission"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(20)
        r2 = p2.add_run(meta["school_mission"])
        r2.italic = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = C.faint

    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after = Pt(20)
    _bottom_border(div, _hex(C.mid), sz=12)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after = Pt(10)
    tr = t.add_run("IT State of the System Report")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = C.text

    if meta.get("respondent_name") or meta.get("respondent_role"):
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.paragraph_format.space_before = Pt(4)
        pf.paragraph_format.space_after = Pt(2)
        _run(pf, "Prepared for", size=10, color=C.faint)
        if meta.get("respondent_name"):
            pn = doc.add_paragraph()
            pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pn.paragraph_format.space_before = Pt(2)
            pn.paragraph_format.space_after = Pt(2)
            _run(pn, meta["respondent_name"], bold=True, size=12)
        if meta.get("respondent_role"):
            pr = doc.add_paragraph()
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr.paragraph_format.space_before = Pt(0)
            pr.paragraph_format.space_after = Pt(4)
            _run(pr, meta["respondent_role"], size=10, color=C.faint)

    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_before = Pt(10)
    pd.paragraph_format.space_after = Pt(2)
    assessment_date = meta.get("assessment_date", "")
    report_date     = meta.get("report_date", date.today().isoformat())
    if assessment_date and assessment_date != report_date:
        _run(pd, f"Assessment conducted: {assessment_date}", size=9, color=C.faint)
        pg = doc.add_paragraph()
        pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pg.paragraph_format.space_before = Pt(0)
        _run(pg, f"Report generated: {report_date}", size=9, color=C.faint)
    else:
        _run(pd, report_date, size=9, color=C.faint)

    # DRAFT watermark — shown when assessment is not yet complete
    if meta.get("is_draft"):
        draft_p = doc.add_paragraph()
        draft_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        draft_p.paragraph_format.space_before = Pt(24)
        draft_p.paragraph_format.space_after = Pt(4)
        dr = draft_p.add_run("DRAFT — ASSESSMENT INCOMPLETE")
        dr.bold = True
        dr.font.size = Pt(14)
        dr.font.color.rgb = C.concern
        draft_sub = doc.add_paragraph()
        draft_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        draft_sub.paragraph_format.space_before = Pt(0)
        _run(draft_sub,
             "Not all sections have been completed. Findings and scores reflect only "
             "the sections assessed so far.",
             size=9, italic=True, color=C.faint)


# ── Condition Summary table ───────────────────────────────────────

def _condition_summary(doc, findings):
    """
    3-column table: Area | Current State | Recommended Change.
    Shows the top 5 findings by severity, giving the reader an
    instant before/after picture of the most important issues.
    Appears inside the Executive Summary, after the Section Scores table.
    """
    # Pick up to 5: urgent first, then concern, then watch
    ranked = sorted(
        [f for f in findings if f.get("severity") in ("urgent", "concern", "watch")],
        key=lambda x: (SEV_ORDER.get(x["severity"], 9), x["finding_id"])
    )[:5]

    if not ranked:
        return

    _h(doc, "Current Conditions vs. Recommended Changes", 2)
    _para(doc,
          "The table below captures the most significant gaps identified in this assessment "
          "and the primary recommended change for each.",
          size=10, sb=4, sa=6, color=C.faint)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"

    for i, txt in enumerate(["Area", "Current Condition", "Recommended Change"]):
        c = tbl.rows[0].cells[i]
        _cell_bg(c, _hex(C.accent))
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C.white

    for idx, f in enumerate(ranked):
        row = tbl.add_row().cells
        fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
        for c in row:
            _cell_bg(c, fill)

        sev_col = SEV_COLOR.get(f["severity"], C.text)

        # Area cell: finding title + severity badge
        row[0].paragraphs[0].clear()
        r0 = row[0].paragraphs[0].add_run(f["title"])
        r0.font.size = Pt(9)
        r0.font.color.rgb = C.text
        r0.bold = True
        p0b = row[0].add_paragraph()
        r0b = p0b.add_run(SEV_LABEL.get(f["severity"], f["severity"]))
        r0b.font.size = Pt(8)
        r0b.font.color.rgb = sev_col
        r0b.bold = True

        # Current condition cell: first ~160 chars of description
        desc = (f.get("description") or "")[:160].rstrip()
        if len(f.get("description") or "") > 160:
            desc += "…"
        row[1].paragraphs[0].clear()
        r1 = row[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(9)
        r1.font.color.rgb = C.text

        # Recommended change cell: first action description
        first_action = (f.get("actions") or [{}])[0]
        act_txt = (first_action.get("description") or "See findings section.")[:160].rstrip()
        if len(first_action.get("description") or "") > 160:
            act_txt += "…"
        row[2].paragraphs[0].clear()
        r2 = row[2].paragraphs[0].add_run(act_txt)
        r2.font.size = Pt(9)
        r2.font.color.rgb = C.text

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Executive Summary ─────────────────────────────────────────────

def _exec_summary(doc, meta, summary, findings, scores, answers=None):
    _page_break(doc)
    _h(doc, "Executive Summary", 1)

    # ── Overall verdict ──────────────────────────────────────────
    verdict = _health_verdict(summary, scores)
    def verdict_builder(cell):
        p = _cp(cell, sb=5, sa=5)
        _run(p, verdict, size=11, bold=False)
    fill = ("FDEDEC" if summary["urgent_count"] >= 2 else
            "FEF9E7" if summary["urgent_count"] >= 1 or summary["concern_count"] >= 3 else
            "EAF4FB")
    _box(doc, fill, verdict_builder)

    # ── Assessment scope callout ─────────────────────────────────
    all_section_ids = set(SECTION_NAMES.keys())
    scored_ids = {str(s["section"]["section_id"]) for s in scores if s.get("max_pts", 0) > 0}
    skipped_ids = all_section_ids - scored_ids

    # DRAFT callout — shown when not all sections are complete
    if meta.get("is_draft") and skipped_ids:
        incomplete_names = sorted(
            [f"Section {sid}: {SECTION_NAMES.get(sid, sid)}" for sid in skipped_ids],
            key=lambda x: int(x.split(":")[0].split()[-1])
        )
        def draft_builder(cell):
            p = _cp(cell, sb=3, sa=2)
            _run(p, "⚠  DRAFT — Assessment Incomplete  ", bold=True, size=10, color=C.concern)
            _run(p, "The following sections have not yet been completed. "
                    "Findings, scores, and the action plan reflect only the sections assessed so far. "
                    "Complete the remaining sections and regenerate this report for a full picture.",
                 size=10)
            for name in incomplete_names:
                pb = _cp(cell, sb=1, sa=1)
                pb.paragraph_format.left_indent = Pt(14)
                _run(pb, f"–  {name}", size=9, color=C.faint)
        _box(doc, "FDEDEC", draft_builder)

    _scope_box(doc, len(scored_ids), skipped_ids)

    # ── Data confidence callout (in body, not just footer) ───────
    caveat = meta.get("confidence_caveat", "")
    if caveat:
        def conf_builder(cell):
            p = _cp(cell, sb=3, sa=3)
            _run(p, "\u26a0  Data confidence note:  ", bold=True, size=10, color=C.concern)
            _run(p, caveat, size=10, italic=True)
        _box(doc, "FEF9E7", conf_builder)

    # ── IT Staffing stub ─────────────────────────────────────────
    if answers:
        _staffing_stub(doc, answers)

    # ── Finding count overview ───────────────────────────────────
    _h(doc, "Assessment Overview", 2)
    p = _para(doc, sb=4, sa=8)
    _run(p, "This assessment identified ")
    _run(p, f"{summary['urgent_count']} urgent", bold=True, color=C.urgent)
    _run(p, ", ")
    _run(p, f"{summary['concern_count']} concern", bold=True, color=C.concern)
    _run(p, ", and ")
    _run(p, f"{summary['watch_count']} watch", bold=True, color=C.watch)
    _run(p, " level findings. ")
    _run(p, "Note: composite finding suppression is not yet enabled in this version — "
           "all findings are listed individually.", italic=True, color=C.faint)

    # ── Overall weighted score ───────────────────────────────────
    # Section weights per scoring framework v0.1 (Sections 1 and 10 excluded).
    SECTION_WEIGHTS = {
        "2": 0.15, "3": 0.12, "4": 0.14, "5": 0.10,
        "6": 0.12, "7": 0.15, "8": 0.12, "9": 0.10,
    }
    if scores:
        weighted_sum = 0.0
        weight_used  = 0.0
        for s in scores:
            sid = str(s["section"]["section_id"])
            w   = SECTION_WEIGHTS.get(sid, 0)
            if w > 0 and s["max_pts"] > 0:
                weighted_sum += s["pct"] * w
                weight_used  += w
        overall_pct = round(weighted_sum / weight_used) if weight_used > 0 else None

        if overall_pct is not None:
            sev_band = ("Healthy"  if overall_pct >= 85 else
                        "Watch"    if overall_pct >= 65 else
                        "Concern"  if overall_pct >= 40 else "Urgent")
            band_col = (C.healthy  if sev_band == "Healthy" else
                        C.watch    if sev_band == "Watch"   else
                        C.concern  if sev_band == "Concern" else C.urgent)

            def score_builder(cell, op=overall_pct, sb=sev_band, bc=band_col):
                p = _cp(cell, sb=4, sa=2)
                _run(p, "Overall Score:  ", bold=True, size=11)
                _run(p, f"{op}%  ", bold=True, size=14, color=bc)
                _run(p, f"({sb})", bold=True, size=11, color=bc)
                p2 = _cp(cell, sb=4, sa=4)
                _run(p2,
                     "This score is a weighted average of the eight scored sections "
                     "(Sections 2–9). Section 1 is school identity context only and "
                     "Section 10 is a calibration input — neither contributes to the "
                     "score. Section weights reflect domain importance: Governance and "
                     "Backup/Recovery are weighted highest at 15% each, reflecting the "
                     "foundational impact of ownership gaps and the irreversibility of "
                     "data loss.",
                     size=9, color=C.faint)
            fill = ("FDEDEC" if sev_band == "Urgent"  else
                    "FEF9E7" if sev_band in ("Watch", "Concern") else "EAF4FB")
            _box(doc, fill, score_builder)

    # ── Score Contribution table ─────────────────────────────────
    # Shows each section's weight, score, and weighted contribution
    # so the reader can see exactly how the overall score was built.
    if scores:
        _h(doc, "Score Breakdown by Section", 2)
        ctbl = doc.add_table(rows=1, cols=4)
        ctbl.style = "Table Grid"
        col_widths = [Inches(2.8), Inches(0.8), Inches(1.2), Inches(1.6)]
        for i, (txt, w) in enumerate(zip(
            ["Section", "Weight", "Section Score", "Weighted Contribution"],
            col_widths,
        )):
            c = ctbl.rows[0].cells[i]
            _cell_bg(c, _hex(C.accent))
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = C.white
            c.width = w

        total_contribution = 0.0
        for idx, s in enumerate(scores):
            sid = str(s["section"]["section_id"])
            row = ctbl.add_row().cells
            fill_r = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill_r)

            is_ctx = sid in ("1", "10") or s["max_pts"] == 0
            w = SECTION_WEIGHTS.get(sid, 0)
            sec_label = f"§{sid}  {s['section']['title']}"

            if is_ctx:
                weight_txt  = "—"
                score_txt   = "Context only — not scored"
                contrib_txt = "—"
                name_col = C.faint
            else:
                pct = s["pct"]
                contrib = round(pct * w, 1)
                total_contribution += pct * w
                weight_txt  = f"{int(w * 100)}%"
                score_txt   = f"{pct}%"
                contrib_txt = f"{contrib:.1f}%"
                name_col = C.text

            for c, txt, clr, bold_flag in [
                (row[0], sec_label,   name_col, False),
                (row[1], weight_txt,  C.faint if is_ctx else C.text, False),
                (row[2], score_txt,   C.faint if is_ctx else C.text, False),
                (row[3], contrib_txt, C.faint if is_ctx else C.text, False),
            ]:
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(txt)
                r.font.size = Pt(9)
                r.font.color.rgb = clr
                r.bold = bold_flag

        # Totals row
        tot_row = ctbl.add_row().cells
        _cell_bg(tot_row[0], "E8ECF0")
        _cell_bg(tot_row[1], "E8ECF0")
        _cell_bg(tot_row[2], "E8ECF0")
        _cell_bg(tot_row[3], "E8ECF0")
        for c, txt, bold_flag in [
            (tot_row[0], "Overall weighted score", True),
            (tot_row[1], "100%", True),
            (tot_row[2], "—",   False),
            (tot_row[3], f"{round(total_contribution)}%", True),
        ]:
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.font.size = Pt(9)
            r.font.color.rgb = C.text
            r.bold = bold_flag

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Section Scores table — includes What's Working column ────
    if scores:
        _h(doc, "Section Scores", 2)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        for i, txt in enumerate(["Section", "Score", "Status", "Answered", "What's Working"]):
            c = tbl.rows[0].cells[i]
            _cell_bg(c, _hex(C.accent))
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = C.white

        for idx, s in enumerate(scores):
            row = tbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill)
            is_ctx = s["max_pts"] == 0
            score_txt = "Context only" if is_ctx else f"{s['earned']} / {s['max_pts']} ({s['pct']}%)"
            sev = s.get("severity", "")
            sev_txt = "—" if is_ctx else SEV_LABEL.get(sev, sev)
            sev_col = C.faint if is_ctx else SEV_COLOR.get(sev, C.text)

            # What's Working: highlight healthy sections
            if is_ctx:
                working_txt = "—"
                working_col = C.faint
            elif sev == "healthy":
                working_txt = "✓ Healthy"
                working_col = C.healthy
            elif sev == "watch":
                working_txt = "Mostly good"
                working_col = C.watch
            else:
                working_txt = "Needs work"
                working_col = C.concern

            for c, txt, clr, bold_flag in [
                (row[0], f"{s['section']['section_id']}. {s['section']['title']}", None, False),
                (row[1], score_txt, None, False),
                (row[2], sev_txt, sev_col, True),
                (row[3], f"{s['answered_count']} answered", None, False),
                (row[4], working_txt, working_col, sev == "healthy"),
            ]:
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(txt)
                r.font.size = Pt(9)
                r.font.color.rgb = clr or C.text
                r.bold = bold_flag

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Current conditions vs. recommended changes ───────────────
    _condition_summary(doc, findings)

    # ── Priority Findings boxes ──────────────────────────────────
    mandatory = {"F2-C01", "F7-C01", "F3-C01", "F3-C02", "F5-C02",
                 "F6-C01", "F6-C02", "F8-C01", "F8-008", "F9-C01"}
    exec_f = [f for f in findings if f["finding_id"] in mandatory or f["severity"] == "urgent"]
    if exec_f:
        _h(doc, "Priority Findings", 2)
        _para(doc,
              "The following findings require immediate attention or leadership awareness.",
              size=11, sb=4, sa=8)
        for f in exec_f:
            fill = "FDEDEC" if f["severity"] == "urgent" else "FEF9E7"
            cost_label, cost_color_hex = _cost_tier(f.get("actions") or [])
            def builder(cell, f=f, cost_label=cost_label, cost_color_hex=cost_color_hex):
                p1 = _cp(cell, sb=2, sa=2)
                _run(p1, f"[{SEV_LABEL.get(f['severity'], f['severity'])}]  ",
                     bold=True, size=10, color=SEV_COLOR.get(f["severity"], C.text))
                _run(p1, f["title"], bold=True, size=11)
                if cost_label:
                    _run(p1, f"   \u25cf {cost_label}",
                         bold=False, size=8,
                         color=RGBColor.from_string(cost_color_hex))
                _cp(cell, f["description"], size=10, sb=2, sa=2)
                if f.get("notes_passthrough"):
                    p3 = _cp(cell, sb=4, sa=2)
                    _run(p3, "IT person noted: ", bold=True, italic=True, size=9, color=C.faint)
                    _run(p3, f["notes_passthrough"], italic=True, size=9)
            _box(doc, fill, builder)


# ── Key Risks ─────────────────────────────────────────────────────

def _key_risks(doc, key_risks, findings):
    if not key_risks:
        return
    _page_break(doc)
    _h(doc, "Key Risks", 1)
    _para(doc,
          "Named risk groups aggregate related findings. Each group has a primary finding "
          "labelled 'Start here' — addressing it has the broadest impact on that risk area.",
          size=11, sb=4, sa=12)

    by_id = {f["finding_id"]: f for f in findings}

    # Sort groups by severity
    sorted_risks = sorted(key_risks, key=lambda g: SEV_ORDER.get(g["severity"], 9))

    for group in sorted_risks:
        fill = ("FDEDEC" if group["severity"] == "urgent" else
                "FEF9E7" if group["severity"] == "concern" else "FDFEFE")

        # Identify primary finding: highest severity then largest effort
        effort_rank = {"L": 0, "M+": 1, "M": 2, "S+": 3, "S": 4}
        group_findings = [by_id[fid] for fid in group["finding_ids"] if fid in by_id]
        primary_fid = None
        if group_findings:
            primary_f = min(
                group_findings,
                key=lambda f: (
                    SEV_ORDER.get(f["severity"], 9),
                    # largest effort first (lowest rank number)
                    min(effort_rank.get(a.get("effort", "M"), 5)
                        for a in f.get("actions", [{}])) if f.get("actions") else 5
                )
            )
            primary_fid = primary_f["finding_id"]

        def builder(cell, g=group, primary=primary_fid):
            p1 = _cp(cell, sb=2, sa=6)
            _run(p1, f"[{SEV_LABEL.get(g['severity'], g['severity'])}]  ",
                 bold=True, size=10, color=SEV_COLOR.get(g["severity"], C.text))
            _run(p1, g["title"], bold=True, size=12)
            for fid in g["finding_ids"]:
                f = by_id.get(fid)
                if not f:
                    continue
                pf = _cp(cell, sb=2, sa=2)
                pf.paragraph_format.left_indent = Pt(12)
                is_primary = (fid == primary)
                _run(pf, f"{fid}  ", bold=True, size=9, color=C.faint)
                _run(pf, f["title"], size=10, bold=is_primary)
                if is_primary:
                    _run(pf, "  ← Start here", bold=True, size=9,
                         color=C.concern)
        _box(doc, fill, builder)


# ── Section-by-Section Findings ───────────────────────────────────

def _section_findings(doc, sections_with_findings, all_scored_sections):
    """
    Render per-section findings. Healthy sections get a ✓ marker.
    Sections not present in sections_with_findings but in all_scored_sections
    are shown as healthy.
    all_scored_sections: list of section_id strings that were scored.
    """
def _section_findings(doc, sections_with_findings, all_scored_sections,
                      finding_contexts=None):
    finding_contexts = finding_contexts or {}
    _page_break(doc)
    _h(doc, "Section-by-Section Findings", 1)
    _para(doc,
          "Findings ordered by severity within each section. "
          "Finding IDs (e.g. F3-004) are used for cross-referencing in the action plan. "
          "Sections with no findings are shown with a healthy marker.",
          size=11, sb=4, sa=12)

    # Build a set of sections that have findings
    sections_map = {s["section_id"]: s for s in sections_with_findings}

    for sid in sorted(all_scored_sections, key=lambda x: int(x)):
        sec_name = SECTION_NAMES.get(sid, f"Section {sid}")
        _h(doc, f"Section {sid}: {sec_name}", 2)

        if sid not in sections_map:
            # Healthy — no findings
            def healthy_builder(cell):
                p = _cp(cell, sb=4, sa=4)
                _run(p, "✓  No findings — all assessed controls in this section are in place.",
                     bold=True, size=10, color=C.healthy)
            _box(doc, "EAFAF1", healthy_builder)
            continue

        sec = sections_map[sid]
        for f in sorted(sec["findings"],
                        key=lambda x: (SEV_ORDER.get(x["severity"], 9), x["finding_id"])):
            sev_hex = _hex(SEV_COLOR.get(f["severity"], C.text))

            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(10)
            tp.paragraph_format.space_after = Pt(4)
            tp.paragraph_format.left_indent = Pt(12)
            _left_border(tp, sev_hex, sz=20)
            _run(tp, f"[{SEV_LABEL.get(f['severity'], f['severity'])}]  ",
                 bold=True, size=10, color=SEV_COLOR.get(f["severity"], C.text))
            _run(tp, f["title"], bold=True, size=11)
            _run(tp, f"  {f['finding_id']}", size=8, color=C.faint)
            cost_label, cost_color_hex = _cost_tier(f.get("actions") or [])
            if cost_label:
                _run(tp, f"   \u25cf {cost_label}", size=8,
                     color=RGBColor.from_string(cost_color_hex))

            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(2)
            dp.paragraph_format.space_after = Pt(6)
            dp.paragraph_format.left_indent = Pt(12)
            _run(dp, f["description"], size=10)

            # IT person passthrough note — blue callout
            if f.get("notes_passthrough"):
                def nb(cell, f=f):
                    p = _cp(cell, sb=2, sa=2)
                    _run(p, "IT person noted: ", bold=True, italic=True, size=9, color=C.mid)
                    _run(p, f["notes_passthrough"], italic=True, size=9)
                _box(doc, "D6EAF8", nb)

            # Plain language note — amber/yellow callout (distinct from action boxes)
            if f.get("plain_language_note"):
                def plb(cell, f=f):
                    p = _cp(cell, sb=2, sa=2)
                    _run(p, "ℹ  Note: ", bold=True, size=9, color=C.watch)
                    _run(p, f["plain_language_note"], size=9, color=C.text)
                _box(doc, "FEF9E7", plb)   # amber, distinct from blue action boxes

            # Actions — green-tinted box, clearly labelled
            if f.get("actions"):
                def action_builder(cell, f=f):
                    ph = _cp(cell, sb=3, sa=2)
                    _run(ph, "Recommended actions:", bold=True, size=10, color=C.healthy)
                    for act in f["actions"]:
                        pa = _cp(cell, sb=2, sa=2)
                        pa.paragraph_format.left_indent = Pt(8)
                        constraint = "⚠  " if act.get("constraint_flag") else "→  "
                        _run(pa, constraint, bold=True, size=9,
                             color=C.concern if act.get("constraint_flag") else C.healthy)
                        _run(pa, act["description"], size=10)
                        horizon_txt = HORIZON_LABEL.get(act["time_horizon"], act["time_horizon"])
                        suffix = f"  [{horizon_txt}"
                        if act.get("constraint_flag"):
                            suffix += "  ⚠ Budget constraint"
                        suffix += "]"
                        _run(pa, suffix, italic=True, size=9, color=C.faint)
                _box(doc, "EAFAF1", action_builder)   # pale green

            # Context note — shown when reviewer has annotated this finding
            ctx = finding_contexts.get(f["finding_id"])
            if ctx:
                def ctx_builder(cell, ctx=ctx):
                    ph = _cp(cell, sb=2, sa=1)
                    _run(ph, "📋  Context note  ", bold=True, size=9, color=C.healthy)
                    _run(ph, f"(added {ctx['added_at'][:10]})", size=8, color=C.faint)
                    pb = _cp(cell, sb=1, sa=2)
                    _run(pb, ctx["note"], size=9, italic=True, color=C.text)
                _box(doc, "EAFAF1", ctx_builder)   # same pale green as actions

        doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Phased Remediation Timeline (replaces bullet Action Plan) ─────

def _timeline_section(doc, timeline):
    """
    The single consolidated action table — the former bullet 'Action Plan'
    section has been removed; this is the only action-planning section.
    """
    _page_break(doc)
    _h(doc, "Action Plan — Phased Timeline", 1)
    _para(doc,
          "All recommended actions grouped by phase, ordered by severity then effort. "
          "Effort ratings: S=½ day · S+=1 day · M=3 days · M+=5 days · L=10 days. "
          "⚠ marks budget or staffing constraints.",
          size=10, sb=4, sa=10, color=C.faint)

    EFFORT_LABEL = {
        "S":  "Quick (½ day)",
        "S+": "Short (1 day)",
        "M":  "Medium (3 days)",
        "M+": "Substantial (5 days)",
        "L":  "Large (10 days)",
    }
    SEV_COLOR_MAP = {"urgent": C.urgent, "concern": C.concern, "watch": C.watch}
    phase_fills = {"1": "1A5276", "2": "1F618D", "3": "2471A3", "4": "2E86C1"}

    for phase in timeline["phases"]:
        acts = phase["actions"]
        if not acts and phase["phase"] in (2, 3, 4):
            continue

        # Phase header
        fill_hex = phase_fills.get(str(phase["phase"]), "2C3E50")
        tbl_h = doc.add_table(rows=1, cols=1)
        tbl_h.style = "Table Grid"
        hcell = tbl_h.rows[0].cells[0]
        _cell_bg(hcell, fill_hex)
        _cell_border(hcell)
        for p in list(hcell.paragraphs):
            p._element.getparent().remove(p._element)
        hp = hcell.add_paragraph()
        hp.paragraph_format.space_before = Pt(4)
        hp.paragraph_format.space_after = Pt(4)
        _run(hp, phase["label"], bold=True, size=12, color=C.white)
        if phase["duration_days"] > 0:
            date_str = f"  ·  {phase['start'].strftime('%d %b %Y')} → {phase['end'].strftime('%d %b %Y')}"
            _run(hp, date_str, size=10, color=C.white)
        else:
            _run(hp, f"  ·  Starting {phase['start'].strftime('%d %b %Y')}", size=10, color=C.white)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        if not acts:
            _para(doc, "No actions in this phase.", size=10, italic=True, color=C.faint, sb=2, sa=8)
        else:
            at = doc.add_table(rows=1, cols=5)
            at.style = "Table Grid"
            for i, txt in enumerate(["Action", "Finding", "Section", "Effort", "Severity"]):
                c = at.rows[0].cells[i]
                _cell_bg(c, "EBF5FB")
                c.paragraphs[0].clear()
                r = c.paragraphs[0].add_run(txt)
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = C.accent

            effort_rank = {"L": 0, "M+": 1, "M": 2, "S+": 3, "S": 4}
            sev_rank    = {"urgent": 0, "concern": 1, "watch": 2}
            sorted_acts = sorted(acts, key=lambda a: (
                sev_rank.get(a.get("severity", "watch"), 9),
                effort_rank.get(a.get("effort", "M"), 5),
            ))

            for idx, act in enumerate(sorted_acts):
                row = at.add_row().cells
                fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
                for c in row:
                    _cell_bg(c, fill)
                sev_col = SEV_COLOR_MAP.get(act.get("severity", "watch"), C.text)
                effort_txt = EFFORT_LABEL.get(act.get("effort", ""), act.get("effort", "—"))
                desc = ("⚠ " if act.get("constraint_flag") else "") + act["description"]
                sec_label = f"§{act.get('section_id', '')}"

                for c, txt, col, bold_flag in [
                    (row[0], desc,                                      C.text,  False),
                    (row[1], act["finding_id"],                         C.faint, False),
                    (row[2], sec_label,                                 C.faint, False),
                    (row[3], effort_txt,                                C.text,  False),
                    (row[4], act.get("severity", "—").upper(),         sev_col, True),
                ]:
                    c.paragraphs[0].clear()
                    r = c.paragraphs[0].add_run(str(txt))
                    r.font.size = Pt(8)
                    r.font.color.rgb = col
                    r.bold = bold_flag

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        if phase.get("rerun_note"):
            def rnb(cell, note=phase["rerun_note"]):
                p = _cp(cell, sb=4, sa=4)
                _run(p, "⚡  ", bold=True, size=10, color=C.mid)
                _run(p, note, size=10, color=C.mid)
            _box(doc, "D6EAF8", rnb)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Strategic / future
    strat = timeline.get("strategic_actions", [])
    if strat:
        _h(doc, "Strategic & Future Actions", 2)
        _para(doc,
              f"{len(strat)} action(s) are flagged as strategic or multi-year initiatives. "
              "Not included in the phased timeline — revisit annually.",
              size=10, sb=4, sa=6, color=C.faint)
        for act in strat:
            p = _para(doc, sb=2, sa=2)
            p.paragraph_format.left_indent = Pt(18)
            _run(p, f"[{act['finding_id']}]  ", size=9, color=C.faint)
            _run(p, act["description"], size=10)


# ── Appendix ──────────────────────────────────────────────────────

def _appendix(doc, suppressed, unknown_log, response_log, amendment_log=None):
    _page_break(doc)
    _h(doc, "Appendix", 1)

    # A. Composite Finding Traceability
    if suppressed:
        _h(doc, "A. Composite Finding Traceability", 2)
        _para(doc,
              "Findings absorbed into composites — their actions are included in the composite finding.",
              size=10, sb=4, sa=8)
        for f in suppressed:
            p = _para(doc, sb=2, sa=2)
            _run(p, f"{f['finding_id']}  ", bold=True, size=9, color=C.faint)
            _run(p, f["title"], size=10)
            _run(p, f"  → {f['suppressed_by']}", italic=True, size=9, color=C.faint)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # B. Unknown Answer Log — prominent framing
    if unknown_log:
        _h(doc, "B. Knowledge Gaps — Unknown Answers", 2)
        def unknown_intro_builder(cell):
            p = _cp(cell, sb=4, sa=2)
            _run(p, f"⚠  {len(unknown_log)} question(s) were answered as 'I don't know.'  ",
                 bold=True, size=10, color=C.concern)
            _run(p, "Each one represents a gap in IT situational awareness — "
                    "something the school does not currently know about its own environment. "
                    "Review each item below and investigate.",
                 size=10)
        _box(doc, "FEF9E7", unknown_intro_builder)

        by_sec = {}
        for u in unknown_log:
            by_sec.setdefault(u["section_id"], []).append(u["question_id"])
        for sid in sorted(by_sec):
            p = _para(doc, sb=2, sa=2)
            p.paragraph_format.left_indent = Pt(18)
            sec_name = SECTION_NAMES.get(sid, f"Section {sid}")
            _run(p, f"Section {sid} — {sec_name}:  ", bold=True, size=10)
            _run(p, ", ".join(by_sec[sid]), size=10)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # C. Full Response Log — with question text
    _h(doc, "C. Full Response Log", 2)
    _para(doc,
          "Complete record of all answers submitted during this assessment, "
          "with question prompts for context.",
          size=10, sb=4, sa=8)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, txt in enumerate(["Q ID", "Question", "Status", "Answer"]):
        c = tbl.rows[0].cells[i]
        _cell_bg(c, _hex(C.accent))
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C.white

    for idx, row_data in enumerate(response_log):
        row = tbl.add_row().cells
        fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
        for c in row:
            _cell_bg(c, fill)
        status_col = (C.concern if row_data["status"] == "unknown" else
                      C.faint if row_data["status"] == "skipped" else C.text)
        prompt = QUESTION_PROMPTS.get(row_data["question_id"], "")
        answer_txt = (row_data["answer"] or f"({row_data['status']})")[:150]
        for c, txt, col in [
            (row[0], row_data["question_id"], C.faint),
            (row[1], prompt,                  C.faint),
            (row[2], row_data["status"],      status_col),
            (row[3], answer_txt,              C.text),
        ]:
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(txt)
            run.font.size = Pt(8)
            run.font.color.rgb = col

    # D. Answer Amendment Log — only if any revisions exist
    if amendment_log:
        _h(doc, "D. Answer Amendment Log", 2)
        _para(doc,
              "The following answers were changed after their section was initially marked complete. "
              "Each row shows the previous value alongside the replacement.",
              size=10, sb=4, sa=8)
        atbl = doc.add_table(rows=1, cols=4)
        atbl.style = "Table Grid"
        for i, txt in enumerate(["Q ID", "Changed At", "Previous Answer", "Revised Answer"]):
            c = atbl.rows[0].cells[i]
            _cell_bg(c, _hex(C.accent))
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = C.white

        for idx, row_data in enumerate(amendment_log):
            row = atbl.add_row().cells
            fill = "FFFFFF" if idx % 2 == 0 else "F8F9FA"
            for c in row:
                _cell_bg(c, fill)
            old_val = str(row_data.get("old_raw_answer") or f"({row_data.get('old_answer_status', '—')})")[:120]
            new_val = str(row_data.get("new_raw_answer") or f"({row_data.get('new_answer_status', '—')})")[:120]
            changed = row_data.get("changed_at", "")[:16].replace("T", " ")
            for c, txt, col in [
                (row[0], row_data["question_id"], C.faint),
                (row[1], changed,                 C.faint),
                (row[2], old_val,                 C.concern),
                (row[3], new_val,                 C.text),
            ]:
                c.paragraphs[0].clear()
                run = c.paragraphs[0].add_run(txt)
                run.font.size = Pt(8)
                run.font.color.rgb = col


# ── Main entry point ──────────────────────────────────────────────

def generate_report(report_data, answers, profile, section_results=None,
                    start_date=None, is_complete=True, finding_contexts=None,
                    amendment_log=None, assessment_date=None):
    school_name     = _get(answers, "1.1") or (profile or {}).get("school_name", "School")
    school_mission  = _get(answers, "1.5")
    respondent_name = _get(answers, "1.7a")
    respondent_role = _get(answers, "1.7b")
    report_date     = date.today().isoformat()
    is_draft        = not is_complete
    finding_contexts = finding_contexts or {}

    confidence = report_data.get("data_confidence", "high")
    caveat_map = {
        "moderate": "Most answers are based on recall rather than documented verification. "
                    "Verify findings before taking action.",
        "low":      "Many answers were estimated or unknown. "
                    "Treat findings as provisional until confirmed.",
        "mixed":    "Data confidence varies by section. "
                    "Verify findings in flagged areas before taking action.",
    }
    confidence_caveat = caveat_map.get(confidence, "")

    meta = {
        "school_name":       school_name,
        "school_mission":    school_mission,
        "respondent_name":   respondent_name,
        "respondent_role":   respondent_role,
        "report_date":       report_date,
        "assessment_date":   assessment_date or report_date,
        "confidence_caveat": confidence_caveat,
        "is_draft":          is_draft,
    }

    findings   = report_data.get("findings", [])
    suppressed = report_data.get("suppressed_findings", [])
    key_risks  = list(report_data.get("key_risk_groups", {}).values())
    by_sev     = report_data.get("by_severity", {})

    summary = {
        "urgent_count":    len(by_sev.get("urgent", [])),
        "concern_count":   len(by_sev.get("concern", [])),
        "watch_count":     len(by_sev.get("watch", [])),
        "suppressed_count": len(suppressed),
        "finding_count":   len(findings),
    }

    # Build sections-with-findings list
    sections_map = {}
    for f in findings:
        sid   = f["section_id"]
        sname = SECTION_NAMES.get(sid, f"Section {sid}")
        sections_map.setdefault(sid, {"section_id": sid, "section_name": sname, "findings": []})["findings"].append(f)
    sections_with_findings = list(sections_map.values())

    # All scored section IDs (from section_results, excluding context-only)
    all_scored_ids = []
    for s in (section_results or []):
        if s.get("max_pts", 0) > 0:
            all_scored_ids.append(str(s["section"]["section_id"]))

    # Unknown answer log
    unknown_log = sorted(
        [{"question_id": qid, "section_id": qid.split(".")[0]}
         for qid, d in answers.items() if d.get("answer_status") == "unknown"],
        key=lambda x: (x["section_id"], x["question_id"])
    )

    # Full response log
    response_log = sorted(
        [{"question_id": qid,
          "section_id":  qid.split(".")[0],
          "status":      d.get("answer_status", "unanswered"),
          "answer":      str(d.get("raw_answer", "")) if d.get("raw_answer") else ""}
         for qid, d in answers.items()],
        key=lambda x: (x["section_id"], x["question_id"])
    )

    # Build phased timeline
    timeline = None
    if start_date:
        from timeline import build_timeline
        from datetime import date as _date
        sd = _date.fromisoformat(start_date) if isinstance(start_date, str) else start_date
        timeline = build_timeline(findings, sd)

    # Add section_id to each timeline action for the §N column
    if timeline:
        fid_to_sid = {f["finding_id"]: f["section_id"] for f in findings}
        for phase in timeline.get("phases", []):
            for act in phase.get("actions", []):
                act["section_id"] = fid_to_sid.get(act.get("finding_id", ""), "")
        for act in timeline.get("strategic_actions", []):
            act["section_id"] = fid_to_sid.get(act.get("finding_id", ""), "")

    # Build document
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

    # Apply heading styles with outlineLevel so TOC works
    _apply_heading_styles(doc)

    _set_hf(doc, school_name, report_date, confidence_caveat,
            assessment_date=meta.get("assessment_date"))
    _cover(doc, meta)
    _toc(doc)
    _exec_summary(doc, meta, summary, findings, section_results or [], answers=answers)
    _key_risks(doc, key_risks, findings)
    _section_findings(doc, sections_with_findings, all_scored_ids,
                      finding_contexts=finding_contexts)
    # Note: bullet Action Plan section removed — Phased Timeline is the single source
    if timeline:
        _timeline_section(doc, timeline)
    _appendix(doc, suppressed, unknown_log, response_log,
              amendment_log=amendment_log)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
