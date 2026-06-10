"""
report_generator_vr.py  —  DOCX report generator for Module 3
Software, Licensing, and Vendor Register
v0.8.2

Report sections:
  1. Cover page
  2. Executive Summary — overall grade, top priorities, key stats
  3. Renewal Risk Register — table sorted by risk level
  4. Category Overview — vendor counts by category
  5. Per-Vendor Report Cards — grade, area scores, findings, strengths
  6. School-Wide Governance Findings (VR2)
  7. Action Plan — by timing bucket (Immediate / Near-Term / Planned)
  8. Phased Timeline (optional, only when start_date supplied)
  9. Appendix — raw answer log per vendor
"""

import io
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ────────────────────────────────────────────────
class C:
    urgent  = RGBColor(0xC0, 0x39, 0x2B)
    concern = RGBColor(0xE6, 0x7E, 0x22)
    watch   = RGBColor(0xF3, 0x9C, 0x12)
    healthy = RGBColor(0x27, 0xAE, 0x60)
    accent  = RGBColor(0x1A, 0x52, 0x76)
    mid     = RGBColor(0x2E, 0x86, 0xC1)
    text    = RGBColor(0x2C, 0x3E, 0x50)
    faint   = RGBColor(0x7F, 0x8C, 0x8D)
    white   = RGBColor(0xFF, 0xFF, 0xFF)
    silver  = RGBColor(0xBD, 0xC3, 0xC7)
    grade_a = RGBColor(0x1E, 0x8B, 0x4C)
    grade_b = RGBColor(0x1A, 0x52, 0x76)
    grade_c = RGBColor(0xB7, 0x77, 0x0D)
    grade_d = RGBColor(0xC0, 0x39, 0x2B)
    grade_f = RGBColor(0x8E, 0x1A, 0x1A)
    draft   = RGBColor(0xC0, 0x39, 0x2B)


SEV_COLOR = {
    "urgent":  C.urgent,
    "concern": C.concern,
    "watch":   C.watch,
    "healthy": C.healthy,
}
SEV_LABEL = {
    "urgent":  "URGENT",
    "concern": "CONCERN",
    "watch":   "WATCH",
    "healthy": "HEALTHY",
}
GRADE_COLOR = {
    "A": C.grade_a, "B": C.grade_b, "C": C.grade_c,
    "D": C.grade_d, "F": C.grade_f,
}
FSEV_COLOR = {
    "critical": RGBColor(0x8E, 0x1A, 0x1A),
    "high":     C.urgent,
    "medium":   C.concern,
    "low":      C.watch,
}
EFFORT_LABEL = {
    "S":  "Quick (½ day)",
    "S+": "Short (1 day)",
    "M":  "Medium (3 days)",
    "M+": "Substantial (5 days)",
    "L":  "Large (10 days)",
}
TIMING_LABEL = {
    "immediate": "Do within 30 days",
    "near_term": "Do within 90 days",
    "planned":   "Schedule this year",
}
TIMING_COLOR = {
    "immediate": RGBColor(0xC0, 0x39, 0x2B),
    "near_term": RGBColor(0xE6, 0x7E, 0x22),
    "planned":   RGBColor(0x27, 0xAE, 0x60),
}
RISK_COLOR = {
    "high":   C.urgent,
    "medium": C.concern,
    "low":    C.healthy,
}

# Question prompts for appendix
QUESTION_PROMPTS = {
    "V.ID.name":          "Vendor / product name",
    "V.ID.category":      "Category",
    "V.ID.status":        "Contract/subscription status",
    "V.ID.owner":         "Named owner at school",
    "V.ID.dept":          "Primary department",
    "V.COST.amount":      "Annual cost",
    "V.COST.known":       "Is cost confirmed and tracked?",
    "V.COST.cycle":       "Billing cycle",
    "V.COST.budget":      "In current budget?",
    "V.RENEW.date":       "Renewal / expiry date",
    "V.RENEW.auto":       "Auto-renews?",
    "V.RENEW.notice":     "Cancellation notice required",
    "V.RENEW.tracked":    "Renewal tracked with reminder?",
    "V.RENEW.signed":     "Signed contract on file?",
    "V.SUPPORT.contact":  "Support contact documented?",
    "V.SUPPORT.escalation": "Escalation path documented?",
    "V.SUPPORT.admin":    "Admin credentials documented?",
    "V.DATA.student":     "Holds student data?",
    "V.DATA.ferpa":       "FERPA/COPPA reviewed?",
    "V.DATA.dpa":         "Signed DPA on file?",
    "V.DATA.staff":       "Holds confidential staff data?",
    "V.USE.active":       "Actively used?",
    "V.USE.value":        "Delivering clear value?",
    "V.USE.notes":        "Notes",
}


# ── Low-level helpers ─────────────────────────────────────────────

def _hex(c):
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def _cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_border(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _left_border(para, hex_color, sz=20):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _run(para, text, bold=False, italic=False, size=11, color=None):
    r = para.add_run(str(text or ""))
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or C.text
    return r


def _h(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    sizes = {1: 16, 2: 13, 3: 11}
    _run(p, text, bold=True, size=sizes.get(level, 11),
         color=color or (C.accent if level == 1 else C.text))
    return p


def _para(doc, text, size=10, color=None, bold=False, italic=False,
          sb=2, sa=2, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    _run(p, text, bold=bold, italic=italic, size=size, color=color or C.text)
    return p


def _cp(cell, sb=4, sa=4):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p


def _box(doc, builder_fn, border_hex="2E86C1", bg_hex=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    if bg_hex:
        _cell_bg(cell, bg_hex)
    _cell_border(cell, color=border_hex, sz="6")
    # Clear default empty paragraph content
    cell.paragraphs[0].clear()
    builder_fn(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _set_hf(doc, school_name, report_date, assessment_date=None):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_label = f"Assessed: {assessment_date}  ·  Generated: {report_date}" \
                     if assessment_date and assessment_date != report_date else report_date
        _run(hp, f"{school_name}  ·  Vendor Register  ·  {date_label}",
             size=8, color=C.faint, italic=True)


# ── Cover page ────────────────────────────────────────────────────

def _cover(doc, school_name, respondent_name, respondent_role, report_date,
           is_draft=False, assessment_date=None):
    if is_draft:
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.paragraph_format.space_before = Pt(6)
        dp.paragraph_format.space_after = Pt(6)
        _run(dp, "DRAFT — ASSESSMENT INCOMPLETE", bold=True, size=10, color=C.draft)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(4)
    _run(title_p, "Software, Licensing &", bold=True, size=26, color=C.accent)

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p2.paragraph_format.space_after = Pt(8)
    _run(title_p2, "Vendor Register", bold=True, size=26, color=C.accent)

    school_p = doc.add_paragraph()
    school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_p.paragraph_format.space_after = Pt(4)
    _run(school_p, school_name, bold=True, size=18, color=C.text)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(4)
    if assessment_date and assessment_date != report_date:
        _run(date_p, f"Assessment conducted: {assessment_date}", size=11, color=C.faint)
        date_p2 = doc.add_paragraph()
        date_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p2.paragraph_format.space_after = Pt(32)
        _run(date_p2, f"Report generated: {report_date}", size=11, color=C.faint)
    else:
        date_p.paragraph_format.space_after = Pt(32)
        _run(date_p, f"Generated {report_date}", size=11, color=C.faint)

    if respondent_name or respondent_role:
        by_p = doc.add_paragraph()
        by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        by_p.paragraph_format.space_after = Pt(4)
        label = respondent_name
        if respondent_role:
            label += f"  ·  {respondent_role}"
        _run(by_p, label, size=11, color=C.faint, italic=True)

    _run(doc.add_paragraph(), "Module 3 — School IT Documentation Engine",
         size=9, italic=True, color=C.silver)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _page_break(doc)


# ── Executive Summary ─────────────────────────────────────────────

def _exec_summary(doc, vr_report, school_name, is_draft=False):
    _h(doc, "Executive Summary")

    s = vr_report.summary

    if is_draft:
        def _draft_warn(cell):
            p = _cp(cell, sb=6, sa=6)
            _run(p, "⚠  DRAFT REPORT — Not all vendor worksheets have been completed. "
                 "Findings and the risk register reflect only the vendors assessed so far.",
                 bold=True, size=10, color=C.draft)
        _box(doc, _draft_warn, border_hex="C0392B", bg_hex="FDECEA")

    # Grade callout
    grade_color = GRADE_COLOR.get(s.overall_grade, C.text)
    _para(doc, f"Overall Grade: {s.overall_grade}  ·  {s.total_vendors} vendor"
          f"{'s' if s.total_vendors != 1 else ''} assessed",
          size=14, bold=True, color=grade_color, sb=8, sa=2)
    _para(doc,
          "The overall grade is a criticality-weighted average. Core systems and vendors "
          "holding student data carry more weight than supplemental tools. "
          "Each vendor's weight is shown on its scorecard below.",
          size=9, color=C.faint, sb=2, sa=8)

    # Quick stats
    stats_tbl = doc.add_table(rows=2, cols=4)
    stats_tbl.style = "Table Grid"
    labels = ["Urgent", "Concern", "Watch", "Healthy"]
    counts = [s.vendors_urgent, s.vendors_concern, s.vendors_watch, s.vendors_healthy]
    colors_sev = [C.urgent, C.concern, C.watch, C.healthy]
    for i, (lbl, cnt, col) in enumerate(zip(labels, counts, colors_sev)):
        hc = stats_tbl.cell(0, i)
        _cell_bg(hc, _hex(col))
        p = _cp(hc, sb=6, sa=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, str(cnt), bold=True, size=16, color=C.white)
        dc = stats_tbl.cell(1, i)
        dp = _cp(dc, sb=2, sa=6)
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(dp, lbl, size=9, color=C.faint)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Key risk stats
    if s.vendors_with_student_data or s.vendors_missing_dpa or s.vendors_auto_renewing_untracked:
        def _risk_stats(cell):
            p = _cp(cell, sb=6, sa=2)
            _run(p, "Key Risk Indicators", bold=True, size=10, color=C.accent)
            items = []
            if s.vendors_with_student_data:
                items.append(f"{s.vendors_with_student_data} vendor{'s' if s.vendors_with_student_data != 1 else ''} hold student data")
            if s.vendors_missing_dpa:
                items.append(f"{s.vendors_missing_dpa} student-data vendor{'s' if s.vendors_missing_dpa != 1 else ''} missing a signed DPA")
            if s.vendors_auto_renewing_untracked:
                items.append(f"{s.vendors_auto_renewing_untracked} auto-renewing contract{'s' if s.vendors_auto_renewing_untracked != 1 else ''} with no tracked reminder")
            for item in items:
                ip = cell.add_paragraph()
                ip.paragraph_format.space_before = Pt(3)
                ip.paragraph_format.space_after = Pt(1)
                ip.paragraph_format.left_indent = Pt(8)
                _run(ip, f"• {item}", size=10, color=C.text)
            cell.add_paragraph().paragraph_format.space_after = Pt(4)
        _box(doc, _risk_stats, border_hex="1A5276", bg_hex="EAF4FB")

    # ── What's Working Well ───────────────────────────────────────
    # Collect distinct strength bullets from healthy/watch vendors
    working_well = []
    seen_strengths = set()
    for vr in vr_report.per_vendor_results:
        if vr.severity in ("healthy", "watch") and vr.strengths:
            for s_txt in vr.strengths:
                if s_txt not in seen_strengths:
                    seen_strengths.add(s_txt)
                    working_well.append((vr.vendor_name, s_txt))
    # Also surface healthy vendors by name as a group summary
    healthy_vendors = [vr.vendor_name for vr in vr_report.per_vendor_results
                       if vr.severity == "healthy"]

    if working_well or healthy_vendors:
        def _well_builder(cell):
            p = _cp(cell, sb=6, sa=2)
            _run(p, "✓  What's Working Well", bold=True, size=10,
                 color=RGBColor(0x1A, 0x6B, 0x3A))
            if healthy_vendors:
                hp = cell.add_paragraph()
                hp.paragraph_format.space_before = Pt(3)
                hp.paragraph_format.space_after = Pt(2)
                hp.paragraph_format.left_indent = Pt(8)
                names = ", ".join(healthy_vendors[:8])
                suffix = f" and {len(healthy_vendors) - 8} more" if len(healthy_vendors) > 8 else ""
                _run(hp, f"Vendors scoring Healthy: {names}{suffix}.",
                     size=9, color=C.text)
            shown = 0
            for vendor_name, s_txt in working_well:
                if shown >= 6:
                    break
                bp = cell.add_paragraph()
                bp.paragraph_format.space_before = Pt(2)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Pt(8)
                _run(bp, f"• {s_txt}", size=9, color=C.text)
                _run(bp, f"  ({vendor_name})", size=8, italic=True,
                     color=RGBColor(0x7F, 0x8C, 0x8D))
                shown += 1
            cell.add_paragraph().paragraph_format.space_after = Pt(4)
        _box(doc, _well_builder, border_hex="1A6B3A", bg_hex="E8F5EC")

    # ── Quick Wins ────────────────────────────────────────────────
    # Surface low-effort findings (S or S+) from any severity level
    quick_wins = [
        f for f in s.top_priorities
        if f.effort in ("S", "S+")
    ]
    # Fall back to near_term/planned with S effort if top_priorities is short
    if len(quick_wins) < 2:
        all_findings = [f for vr in vr_report.per_vendor_results for f in vr.findings]
        all_findings += vr_report.school_wide_results
        for f in all_findings:
            if f.effort in ("S", "S+") and f not in quick_wins:
                quick_wins.append(f)
                if len(quick_wins) >= 5:
                    break

    if quick_wins:
        def _qw_builder(cell):
            p = _cp(cell, sb=6, sa=2)
            _run(p, "→  Quick Wins Available  (Low Effort)",
                 bold=True, size=10, color=RGBColor(0xB7, 0x77, 0x0D))
            _run(cell.add_paragraph(), "", size=2)  # small gap
            for f in quick_wins[:5]:
                qp = cell.add_paragraph()
                qp.paragraph_format.space_before = Pt(3)
                qp.paragraph_format.space_after = Pt(1)
                qp.paragraph_format.left_indent = Pt(8)
                _run(qp, f"• ", bold=True, size=9,
                     color=RGBColor(0xB7, 0x77, 0x0D))
                _run(qp, f.title, bold=True, size=9, color=C.text)
                if f.vendor_name:
                    _run(qp, f"  ({f.vendor_name})", size=8, italic=True, color=C.faint)
                ap = cell.add_paragraph()
                ap.paragraph_format.space_before = Pt(1)
                ap.paragraph_format.space_after = Pt(3)
                ap.paragraph_format.left_indent = Pt(16)
                _run(ap, f.action, size=8, italic=True, color=C.faint)
            cell.add_paragraph().paragraph_format.space_after = Pt(4)
        _box(doc, _qw_builder, border_hex="B7770D", bg_hex="FEF9EF")

    # ── Investigate Before Next Renewal ──────────────────────────
    # Critical/urgent vendors that need attention — compact action table
    urgent_vendors = [
        vr for vr in vr_report.per_vendor_results
        if vr.severity in ("urgent",) and vr.findings
    ]
    concern_vendors = [
        vr for vr in vr_report.per_vendor_results
        if vr.severity == "concern" and vr.findings
    ]
    investigate = (urgent_vendors + concern_vendors)[:8]

    if investigate:
        _h(doc, "Investigate Before Next Renewal", 2)
        _para(doc,
              "The vendors below have one or more critical or high findings. "
              "Review and act before their next renewal date.",
              size=9, color=C.faint, sb=2, sa=6)

        inv_tbl = doc.add_table(rows=1, cols=4)
        inv_tbl.style = "Table Grid"
        inv_headers = ["Vendor", "Top Issue", "Action", "Priority"]
        inv_widths = [1.3, 2.0, 2.5, 0.9]
        for i, (hdr, w) in enumerate(zip(inv_headers, inv_widths)):
            cell = inv_tbl.cell(0, i)
            _cell_bg(cell, "2C3E50")
            _cell_border(cell, "2C3E50")
            cell.width = Inches(w)
            p = _cp(cell, sb=5, sa=5)
            _run(p, hdr, bold=True, size=9, color=C.white)

        for vr in investigate:
            # Pick the most severe finding
            top_f = sorted(
                vr.findings,
                key=lambda f: {"critical": 0, "high": 1, "medium": 2, "low": 3}.get(f.severity, 9)
            )[0]
            sev_col = FSEV_COLOR.get(top_f.severity, C.text)
            row = inv_tbl.add_row()
            vals = [vr.vendor_name, top_f.title, top_f.action, top_f.severity.upper()]
            for i, val in enumerate(vals):
                cell = row.cells[i]
                _cell_border(cell, "CCCCCC", "2")
                p = _cp(cell, sb=4, sa=4)
                if i == 3:
                    _run(p, val, bold=True, size=9, color=sev_col)
                elif i == 0:
                    _run(p, val, bold=True, size=9, color=C.text)
                else:
                    _run(p, val, size=9, color=C.text)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Top priorities (numbered action list) ─────────────────────
    if s.top_priorities:
        _h(doc, "Top Priorities", 2)
        for idx, f in enumerate(s.top_priorities[:6], start=1):
            sev_col = FSEV_COLOR.get(f.severity, C.text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(6)
            _left_border(p, _hex(sev_col), sz=16)
            _run(p, f"{idx}.  ", bold=True, size=11,
                 color=RGBColor(0xBD, 0xC3, 0xC7))
            _run(p, f"[{f.severity.upper()}]  ", bold=True, size=10, color=sev_col)
            _run(p, f.title, bold=True, size=10, color=C.text)
            if f.vendor_name:
                _run(p, f"  —  {f.vendor_name}", size=9, italic=True, color=C.faint)
            np = doc.add_paragraph()
            np.paragraph_format.left_indent = Pt(22)
            np.paragraph_format.space_before = Pt(1)
            np.paragraph_format.space_after = Pt(2)
            _run(np, f.action, size=9, color=C.faint, italic=True)
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.left_indent = Pt(22)
            meta_p.paragraph_format.space_before = Pt(1)
            meta_p.paragraph_format.space_after = Pt(6)
            _run(meta_p,
                 f"Owner: {f.owner}  ·  Timing: {TIMING_LABEL.get(f.timing, f.timing)}  ·  Effort: {EFFORT_LABEL.get(f.effort, f.effort)}",
                 size=8, color=C.faint)

    _page_break(doc)


# ── Renewal Risk Register ─────────────────────────────────────────

def _renewal_register(doc, vr_report, answers=None):
    _h(doc, "Renewal Risk Register")
    _para(doc,
          "Vendors sorted by renewal risk. High-risk items should be actioned before "
          "the next renewal date. Auto-renewing contracts without tracked reminders "
          "commit budget without a deliberate decision.",
          size=10, color=C.faint, sb=2, sa=8)

    if not vr_report.renewal_risk_register:
        _para(doc, "No renewal data recorded yet. Complete vendor worksheets to populate this register.",
              size=10, color=C.faint)
        _page_break(doc)
        return

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    headers = ["Vendor", "Category", "Renewal Date", "Auto-Renews", "Risk", "Reason"]
    widths  = [1.4, 1.3, 1.0, 0.8, 0.6, 2.2]
    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = tbl.cell(0, i)
        _cell_bg(cell, "2C3E50")
        _cell_border(cell, "2C3E50")
        cell.width = Inches(w)
        p = _cp(cell, sb=5, sa=5)
        _run(p, hdr, bold=True, size=9, color=C.white)

    for r in vr_report.renewal_risk_register:
        row = tbl.add_row()
        risk_col = RISK_COLOR.get(r.risk_level, C.text)
        vals = [r.vendor_name, r.category, r.renewal_date,
                "Yes" if r.auto_renews else "No",
                r.risk_level.upper(), r.risk_reason]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            _cell_border(cell, "CCCCCC", "2")
            p = _cp(cell, sb=4, sa=4)
            if i == 4:
                _run(p, val, bold=True, size=9, color=risk_col)
            elif i == 0:
                _run(p, val, bold=True, size=9, color=C.text)
            else:
                _run(p, val, size=9, color=C.text)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── VR2.10 — Biggest unresolved vendor concern (notes passthrough) ──
    # Quote the respondent's free-text concern directly in the register
    # so it surfaces alongside renewal risk data rather than only in the appendix.
    if answers:
        vr210 = answers.get("VR2.10", {})
        vr210_text = (vr210.get("raw_answer") or "") if isinstance(vr210, dict) else ""
        if vr210_text and str(vr210_text).strip():
            concern_text = str(vr210_text).strip()
            def vr210_builder(cell, ct=concern_text):
                p = _cp(cell, sb=4, sa=2)
                _run(p, "Respondent's biggest unresolved vendor concern  ",
                     bold=True, size=10, color=C.accent)
                p2 = _cp(cell, sb=2, sa=4)
                _run(p2, f'"{ct}"', size=10, italic=True, color=C.text)
            _box(doc, vr210_builder, border_hex="1A5276", bg_hex="EAF4FB")
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _page_break(doc)


# ── Category Overview ─────────────────────────────────────────────

def _category_overview(doc, vr_report):
    if not vr_report.summary.category_breakdown:
        return
    _h(doc, "Vendor Category Breakdown")
    _para(doc, "Distribution of assessed vendors by category.",
          size=10, color=C.faint, sb=2, sa=8)

    breakdown = sorted(vr_report.summary.category_breakdown.items(),
                       key=lambda x: -x[1])
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    for i, hdr in enumerate(["Category", "Count"]):
        cell = tbl.cell(0, i)
        _cell_bg(cell, "2C3E50")
        _cell_border(cell, "2C3E50")
        p = _cp(cell, sb=5, sa=5)
        _run(p, hdr, bold=True, size=9, color=C.white)

    for cat, cnt in breakdown:
        row = tbl.add_row()
        for i, val in enumerate([cat, str(cnt)]):
            cell = row.cells[i]
            _cell_border(cell, "CCCCCC", "2")
            p = _cp(cell, sb=4, sa=4)
            _run(p, val, size=9, color=C.text)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Per-Vendor Findings ───────────────────────────────────────────

def _per_vendor_findings(doc, vr_report, finding_contexts=None):
    finding_contexts = finding_contexts or {}
    _page_break(doc)
    _h(doc, "Per-Vendor Findings")
    _para(doc,
          "One section per vendor. Findings are ordered by severity. "
          "Effort ratings: S=½ day · S+=1 day · M=3 days · M+=5 days · L=10 days.",
          size=10, color=C.faint, sb=4, sa=10)

    for result in vr_report.per_vendor_results:
        _h(doc, result.vendor_name, 2)

        # Score line
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(4)
        sp.paragraph_format.left_indent = Pt(6)
        grade_col = GRADE_COLOR.get(result.grade_label, C.text)
        sev_col = SEV_COLOR.get(result.severity, C.text)
        _run(sp, "Grade: ", size=10, color=C.faint)
        _run(sp, result.grade_label, bold=True, size=12, color=grade_col)
        _run(sp, f"   Score: {result.score_pct}%   Category: {result.category}",
             size=10, color=C.faint)
        _run(sp, "   Status: ", size=10, color=C.faint)
        _run(sp, SEV_LABEL.get(result.severity, result.severity),
             bold=True, size=10, color=sev_col)
        # Weight label — only show when above baseline so routine vendors aren't cluttered
        wm = getattr(result, "weight_multiplier", 1)
        if wm > 1:
            weight_labels = {2: "Staff data — 2× weight",
                             3: "Core or student data — 3× weight",
                             4: "Core + student data — 4× weight"}
            _run(sp, f"   ·  {weight_labels.get(wm, f'{wm}× weight')}",
                 size=9, italic=True, color=C.concern)

        if result.holds_student_data or result.holds_staff_data:
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(2)
            dp.paragraph_format.space_after = Pt(4)
            dp.paragraph_format.left_indent = Pt(6)
            labels = []
            if result.holds_student_data:
                labels.append("student data")
            if result.holds_staff_data:
                labels.append("staff data")
            _run(dp, "Sensitive data: ", size=9, color=C.faint)
            _run(dp, " and ".join(labels), size=9, italic=True, color=C.concern)

        if result.renewal_date:
            rp = doc.add_paragraph()
            rp.paragraph_format.left_indent = Pt(6)
            rp.paragraph_format.space_before = Pt(1)
            rp.paragraph_format.space_after = Pt(6)
            _run(rp, "Renewal: ", size=9, color=C.faint)
            _run(rp, result.renewal_date, size=9, color=C.text)
            if result.auto_renews:
                _run(rp, "  (auto-renews)", size=9, italic=True, color=C.concern)

        if result.area_scores:
            _area_score_table(doc, result.area_scores)

        if not result.findings:
            def _healthy_builder(cell):
                p = _cp(cell, sb=4, sa=2)
                _run(p, "✓  All assessed controls appear to be in place.",
                     bold=True, size=10, color=C.healthy)
                if result.strengths:
                    for s in result.strengths:
                        sp = cell.add_paragraph()
                        sp.paragraph_format.left_indent = Pt(12)
                        sp.paragraph_format.space_before = Pt(2)
                        sp.paragraph_format.space_after = Pt(1)
                        _run(sp, f"• {s}", size=9, color=C.faint)
                cell.add_paragraph().paragraph_format.space_after = Pt(2)
            _box(doc, _healthy_builder, border_hex="27AE60", bg_hex="EAFAF1")
        else:
            for f in sorted(result.findings,
                            key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3}.get(x.severity, 9)):
                _finding_box(doc, f, finding_contexts, section_id=result.section_id)


def _area_score_table(doc, area_scores):
    tbl = doc.add_table(rows=1, cols=len(area_scores))
    tbl.style = "Table Grid"
    for i, (area, (earned, max_pts)) in enumerate(area_scores.items()):
        cell = tbl.cell(0, i)
        pct = round(earned / max_pts * 100) if max_pts > 0 else 0
        if pct >= 85:
            bg = "27AE60"
        elif pct >= 65:
            bg = "F39C12"
        elif pct >= 40:
            bg = "E67E22"
        else:
            bg = "C0392B"
        _cell_bg(cell, bg)
        _cell_border(cell, "FFFFFF", "4")
        p = _cp(cell, sb=4, sa=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, area, bold=True, size=7, color=C.white)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(4)
        _run(p2, f"{pct}%", size=9, bold=True, color=C.white)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _finding_box(doc, f, finding_contexts=None, section_id=None):
    finding_contexts = finding_contexts or {}
    sev_col = FSEV_COLOR.get(f.severity, C.text)
    sev_hex = _hex(sev_col)

    # P2-H2: Use stable rule_id as canonical context key.
    # Format: "{section_id}:{rule_id}". Fall back to area prefix only when
    # rule_id is absent (forward-compatibility guard).
    if section_id:
        if f.rule_id:
            fid_key = f"{section_id}:{f.rule_id}"
        else:
            fid_key = f"{section_id}:{f.area[:3].upper()}"
    else:
        fid_key = None
    ctx = finding_contexts.get(fid_key) if fid_key else None

    def _builder(cell):
        p = _cp(cell, sb=4, sa=2)
        _run(p, f"[{f.severity.upper()}]  ", bold=True, size=10, color=sev_col)
        _run(p, f.title, bold=True, size=10, color=C.text)

        meta = cell.add_paragraph()
        meta.paragraph_format.space_before = Pt(1)
        meta.paragraph_format.space_after = Pt(3)
        meta.paragraph_format.left_indent = Pt(6)
        timing = getattr(f, "timing", "planned")
        _run(meta, f"{f.area}  ·  {TIMING_LABEL.get(timing, timing)}  ·  "
             f"{EFFORT_LABEL.get(f.effort, f.effort)}  ·  {f.owner}",
             size=8, color=C.faint, italic=True)

        dp = cell.add_paragraph()
        dp.paragraph_format.space_before = Pt(4)
        dp.paragraph_format.space_after = Pt(4)
        dp.paragraph_format.left_indent = Pt(6)
        _run(dp, f.detail, size=9, color=C.text)

        ap = cell.add_paragraph()
        ap.paragraph_format.space_before = Pt(2)
        ap.paragraph_format.space_after = Pt(6)
        ap.paragraph_format.left_indent = Pt(6)
        _run(ap, "→ ", bold=True, size=9, color=C.accent)
        _run(ap, f.action, size=9, color=C.text, italic=True)

        # Context note — rendered when a reviewer has annotated this finding
        if ctx:
            ph = cell.add_paragraph()
            ph.paragraph_format.space_before = Pt(6)
            ph.paragraph_format.space_after = Pt(1)
            ph.paragraph_format.left_indent = Pt(6)
            _run(ph, "📋  Context note  ", bold=True, size=9, color=C.healthy)
            _run(ph, f"(added {ctx['added_at'][:10]})", size=8, color=C.faint)
            pb = cell.add_paragraph()
            pb.paragraph_format.space_before = Pt(1)
            pb.paragraph_format.space_after = Pt(4)
            pb.paragraph_format.left_indent = Pt(6)
            _run(pb, ctx["note"], size=9, italic=True, color=C.text)

    _box(doc, _builder, border_hex=sev_hex)


# ── School-Wide Governance Findings ──────────────────────────────

def _school_wide_findings(doc, vr_report, finding_contexts=None):
    finding_contexts = finding_contexts or {}
    if not vr_report.school_wide_results:
        return
    _page_break(doc)
    _h(doc, "School-Wide Vendor Governance Findings")
    _para(doc,
          "These findings relate to the policies, processes, and ownership structures "
          "that govern the school's vendor portfolio as a whole — not any individual vendor. "
          "They are drawn from Section VR2 (School-Wide Vendor Governance).",
          size=10, color=C.faint, sb=4, sa=10)
    for f in sorted(vr_report.school_wide_results,
                    key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3}.get(x.severity, 9)):
        _finding_box(doc, f, finding_contexts, section_id="VR2")


# ── Action Plan ───────────────────────────────────────────────────

def _action_plan(doc, vr_report):
    _page_break(doc)
    _h(doc, "Action Plan")
    _para(doc, "All findings grouped by timing bucket. "
          "Assign a named owner to each action before the next renewal cycle.",
          size=10, color=C.faint, sb=4, sa=8)

    all_findings = [f for r in vr_report.per_vendor_results
                    for f in r.findings] + vr_report.school_wide_results

    buckets = {
        "immediate": ("Immediate — Do within 30 days", C.urgent),
        "near_term": ("Near-Term — Do within 90 days", C.concern),
        "planned":   ("Planned — Schedule this year",  C.healthy),
    }

    for timing_key, (timing_label, timing_col) in buckets.items():
        bucket = [f for f in all_findings
                  if getattr(f, "timing", "planned") == timing_key]
        if not bucket:
            continue

        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(10)
        hp.paragraph_format.space_after  = Pt(4)
        _run(hp, timing_label, bold=True, size=12, color=timing_col)

        for f in sorted(bucket, key=lambda x: {"critical": 0, "high": 1,
                                                 "medium": 2, "low": 3}.get(x.severity, 9)):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Pt(6)
            sev_col = FSEV_COLOR.get(f.severity, C.text)
            _left_border(p, _hex(sev_col), sz=14)
            _run(p, f"[{f.severity.upper()}]  ", bold=True, size=9, color=sev_col)
            _run(p, f.title, bold=True, size=9, color=C.text)
            if getattr(f, "vendor_name", None):
                _run(p, f"  ({f.vendor_name})", size=8, color=C.faint, italic=True)

            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent  = Pt(14)
            ap.paragraph_format.space_before = Pt(1)
            ap.paragraph_format.space_after  = Pt(6)
            _run(ap, f"→ {f.action}  ", size=9, color=C.text, italic=True)
            _run(ap, f"[{EFFORT_LABEL.get(f.effort, f.effort)}  ·  {f.owner}]",
                 size=8, color=C.faint)


# ── Phased Timeline ───────────────────────────────────────────────

def _vr_sev_to_timeline(sev):
    return {"critical": "urgent", "high": "urgent",
            "medium": "concern", "low": "watch"}.get(sev, "watch")


def _vr_timing_to_horizon(timing):
    return {"immediate": "immediate", "near_term": "next_90_days",
            "planned": "this_year"}.get(timing, "next_90_days")


def _timeline_section(doc, timeline):
    from report_generator_dg import _dg_timeline_section as _tl
    _tl(doc, timeline)


# ── Appendix ──────────────────────────────────────────────────────

def _appendix(doc, vr_report, answers, vendor_names, generated_section_ids,
              amendment_log=None):
    _page_break(doc)
    _h(doc, "Appendix A — Raw Answer Log")
    _para(doc,
          "Complete record of all answers recorded in per-vendor worksheets. "
          "Blank cells indicate the question was not answered or not reached.",
          size=10, color=C.faint, sb=4, sa=10)

    for name, sid in zip(vendor_names, generated_section_ids):
        _h(doc, name, 2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        for i, hdr in enumerate(["Question ID", "Question", "Answer"]):
            cell = tbl.cell(0, i)
            _cell_bg(cell, "2C3E50")
            _cell_border(cell, "2C3E50")
            p = _cp(cell, sb=5, sa=5)
            _run(p, hdr, bold=True, size=8, color=C.white)

        for template_qid, prompt in QUESTION_PROMPTS.items():
            full_qid = f"{sid}_{template_qid}"
            rec = answers.get(full_qid, {})
            raw = rec.get("raw_answer", "") if isinstance(rec, dict) else ""
            if isinstance(raw, list):
                raw = "\n".join(str(x) for x in raw)
            elif raw is None:
                raw = ""

            row = tbl.add_row()
            vals = [template_qid, prompt, str(raw)]
            for i, val in enumerate(vals):
                cell = row.cells[i]
                _cell_border(cell, "CCCCCC", "2")
                p = _cp(cell, sb=3, sa=3)
                _run(p, val, size=8, color=C.text if i < 2 else C.faint)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # VR2 School-Wide Governance appendix
    _h(doc, "Section VR2 — School-Wide Vendor Governance", 2)
    vr2_questions = [
        ("VR2.1",  "Software approval process"),
        ("VR2.2",  "Contract signing authority"),
        ("VR2.3",  "Spend threshold policy"),
        ("VR2.4",  "Master vendor register existed before this audit?"),
        ("VR2.5",  "Centralized password manager for vendor accounts?"),
        ("VR2.6",  "Renewal tracking shared between IT and Business Office?"),
        ("VR2.7",  "Vendor account offboarding covered?"),
        ("VR2.8",  "DPAs centrally tracked?"),
        ("VR2.9",  "Annual vendor review conducted?"),
        ("VR2.10", "Biggest unresolved vendor concern"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, hdr in enumerate(["Question ID", "Question", "Answer"]):
        cell = tbl.cell(0, i)
        _cell_bg(cell, "2C3E50")
        _cell_border(cell, "2C3E50")
        p = _cp(cell, sb=5, sa=5)
        _run(p, hdr, bold=True, size=8, color=C.white)
    for qid, prompt in vr2_questions:
        rec = answers.get(qid, {})
        raw = rec.get("raw_answer", "") if isinstance(rec, dict) else ""
        if raw is None:
            raw = ""
        row = tbl.add_row()
        for i, val in enumerate([qid, prompt, str(raw)]):
            cell = row.cells[i]
            _cell_border(cell, "CCCCCC", "2")
            p = _cp(cell, sb=3, sa=3)
            _run(p, val, size=8, color=C.text if i < 2 else C.faint)

    if amendment_log:
        _page_break(doc)
        _h(doc, "Appendix B — Answer Amendment Log")
        _para(doc,
              "The following answers were revised after their section was first marked complete.",
              size=10, color=C.faint, sb=4, sa=8)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, hdr in enumerate(["Question ID", "Previous Answer", "Revised Answer", "Changed"]):
            cell = tbl.cell(0, i)
            _cell_bg(cell, "2C3E50")
            _cell_border(cell, "2C3E50")
            p = _cp(cell, sb=5, sa=5)
            _run(p, hdr, bold=True, size=8, color=C.white)
        for entry in amendment_log:
            row = tbl.add_row()
            vals = [
                entry.get("question_id", ""),
                str(entry.get("old_raw_answer", "")),
                str(entry.get("new_raw_answer", "")),
                str(entry.get("changed_at", ""))[:16],
            ]
            for i, val in enumerate(vals):
                cell = row.cells[i]
                _cell_border(cell, "CCCCCC", "2")
                p = _cp(cell, sb=3, sa=3)
                _run(p, val, size=8, color=C.faint)


# ── Main entry point ──────────────────────────────────────────────

def generate_vr_report(vr_report_obj, answers, profile, vendor_names,
                       generated_section_ids, start_date=None,
                       is_complete=False, finding_contexts=None,
                       amendment_log=None, assessment_date=None):
    """
    Generate the Vendor Register DOCX report.
    Returns bytes.
    """
    school_name = (profile or {}).get("school_name", "School")
    report_date = date.today().isoformat()
    is_draft = not is_complete
    finding_contexts = finding_contexts or {}

    def _get(qid):
        d = answers.get(qid)
        if not d:
            return ""
        r = d.get("raw_answer")
        return ", ".join(str(x) for x in r) if isinstance(r, list) else (str(r) if r else "")

    respondent_name = _get("VR1.1")
    respondent_role = _get("VR1.1b")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

    _set_hf(doc, school_name, report_date, assessment_date=assessment_date)
    _cover(doc, school_name, respondent_name, respondent_role, report_date,
           is_draft=is_draft, assessment_date=assessment_date)
    _exec_summary(doc, vr_report_obj, school_name, is_draft=is_draft)
    _renewal_register(doc, vr_report_obj, answers=answers)
    _category_overview(doc, vr_report_obj)
    _per_vendor_findings(doc, vr_report_obj, finding_contexts=finding_contexts)
    _school_wide_findings(doc, vr_report_obj, finding_contexts=finding_contexts)
    _action_plan(doc, vr_report_obj)

    if start_date:
        from timeline import build_timeline
        from datetime import date as _date
        sd = _date.fromisoformat(start_date) if isinstance(start_date, str) else start_date
        flat_findings = []
        for result in vr_report_obj.per_vendor_results:
            for f in result.findings:
                flat_findings.append({
                    "finding_id":   f"{result.section_id}:{f.rule_id}" if f.rule_id else f"{result.section_id}:{f.area[:3].upper()}",
                    "title":        f.title,
                    "severity":     _vr_sev_to_timeline(f.severity),
                    "section_id":   result.section_id,
                    "actions": [{
                        "description":    f.action,
                        "effort":         f.effort or "M",
                        "time_horizon":   _vr_timing_to_horizon(getattr(f, "timing", "planned")),
                        "constraint_flag": False,
                    }],
                })
        for f in vr_report_obj.school_wide_results:
            flat_findings.append({
                "finding_id":   f"VR2:{f.rule_id}" if f.rule_id else f"VR2:{f.area[:3].upper()}",
                "title":        f.title,
                "severity":     _vr_sev_to_timeline(f.severity),
                "section_id":   "VR2",
                "actions": [{
                    "description":    f.action,
                    "effort":         f.effort or "M",
                    "time_horizon":   _vr_timing_to_horizon(getattr(f, "timing", "planned")),
                    "constraint_flag": False,
                }],
            })
        if flat_findings:
            from datetime import date as _date2
            timeline = build_timeline(flat_findings, sd)
            from report_generator_dg import _dg_timeline_section
            _dg_timeline_section(doc, timeline)

    _appendix(doc, vr_report_obj, answers, vendor_names, generated_section_ids,
              amendment_log=amendment_log)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
